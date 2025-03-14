import customtkinter as ctk
from tkinter import messagebox, ttk, filedialog, font
import pymysql
import pandas as pd
import os, bcrypt, re
from PIL import Image
import tkinter as tk
import datetime, mysql.connector, openpyxl
from tkcalendar import DateEntry
from num2words import num2words 
from dateutil.relativedelta import relativedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx2pdf import convert
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Border, Side
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

#Colores diseño nuevo
COLOR_FONDO = "white"  
COLOR_FRAME = "white"  
COLOR_BOTON = "#058A15"  
COLOR_BOTON_HOVER = "green"  
COLOR_TEXTO = "black"
pantalla=None

# Clase para manejar la base de datos
class DataBase:
    def __init__(self):
        self.connection = pymysql.connect(
            host="localhost",
            user="root",
            password="",
            db="contrataciones"
        )
        self.cursor = self.connection.cursor()

    def ejecutar_consulta(self, consulta, valores=None):
        try:
            if valores:
                self.cursor.execute(consulta, valores)
            else:
                self.cursor.execute(consulta)
            self.connection.commit()
        except Exception as e:
            messagebox.showerror("Error en consulta", f"{e}")
            raise  # Lanzar la excepción para detener la ejecución

    def obtener_datos(self, consulta):
        try:
            self.cursor.execute(consulta)
            return self.cursor.fetchall()
        except Exception as e:
            messagebox.showerror("Error en consulta", f"{e}")
            return []

    def __del__(self):
        self.cursor.close()
        self.connection.close()

# Variables globales
usuario_iniciado = True  # Cambia a False si el usuario aún no ha iniciado sesión
menu_expandido = False
db = DataBase()  # Instancia de la base de datos

def abrir_pantalla(contenido_frame, titulo):
    for widget in contenido_frame.winfo_children():
        widget.destroy()
    
    fuente = ctk.CTkFont(family="Zurich.ttf", weight="bold", size=25)
    label_titulo = ctk.CTkLabel(contenido_frame, text=titulo, font=fuente, text_color="#446344")
    label_titulo.pack(pady=20)
    
    funciones_pantalla = {
        "Pantalla EPS y ARL": lambda frame: mostrar_epsarl(frame, db),
        "Pantalla Bancos": mostrar_bancos,
        "Pantalla Ciudad y Departamento": lambda frame:mostrar_ciudades_departamentos(frame,db),
        "Pantalla Cargo": lambda frame: mostrar_cargo(frame,db),
        "Pantalla Tipo de contrato": mostrar_tipodecontrato,
        "Pantalla Áreas": mostrar_areas,
        "Pantalla Certificado": mostrar_certificado,
        "Pantalla Contratistas": mostrar_Contratista,
        "Pantalla Usuarios": lambda frame:mostrar_usuario(frame,db)
    }
    
    if titulo in funciones_pantalla:
        funciones_pantalla[titulo](contenido_frame)

############################################# Pantalla EPS y ARL ######################################################
def mostrar_epsarl(frame, db):
    nombre_eps = ctk.StringVar()
    accion_eps = ctk.StringVar(value="Agregar")
    eps_id = None

    nombre_arl = ctk.StringVar()
    accion_arl = ctk.StringVar(value="Agregar")
    arl_id = None

    def seleccionar_arl(event):
        nonlocal arl_id
        seleccion = tvARL.selection()
        if seleccion:
            valores = tvARL.item(seleccion[0], "values")
            if valores:
                arl_id = valores[0]
                nombre_arl.set(valores[1])  # Mostrar nombre en el campo de texto

    def seleccionar_eps(event):
        nonlocal eps_id
        seleccionado = tvEPS.selection()
        if seleccionado:
            valores = tvEPS.item(seleccionado[0], "values")
            if valores:
                eps_id = valores[0]
                nombre_eps.set(valores[1])  # Mostrar nombre en el campo de texto

    def llenar_tabla_arl():
        tvARL.delete(*tvARL.get_children())
        for fila in db.obtener_datos("SELECT id, nombreARL FROM arl"):
            tvARL.insert("", "end", values=fila)

    def llenar_tabla_eps():
        tvEPS.delete(*tvEPS.get_children())
        for fila in db.obtener_datos("SELECT id, nombreEPS FROM eps"):
            tvEPS.insert("", "end", values=fila)

    def ejecutar_accion_arl():
        accion = accion_arl.get().strip()
        
        if accion == "Agregar" and nombre_arl.get().strip():
            db.ejecutar_consulta("INSERT INTO arl (nombreARL) VALUES (%s)", (nombre_arl.get(),))
            llenar_tabla_arl()
        
        elif accion == "Modificar" and arl_id and nombre_arl.get().strip():
            db.ejecutar_consulta("UPDATE arl SET nombreARL=%s WHERE id=%s", (nombre_arl.get(), arl_id))
            llenar_tabla_arl()
        
        elif accion == "Eliminar" and arl_id:
            db.ejecutar_consulta("DELETE FROM arl WHERE id=%s", (arl_id,))
            llenar_tabla_arl()

    def ejecutar_accion_eps():
        texto = nombre_eps.get().strip()
        if not texto:
            return

        if accion_eps.get() == "Agregar":
            db.ejecutar_consulta("INSERT INTO eps (nombreEPS) VALUES (%s)", (texto,))
        elif accion_eps.get() == "Modificar" and eps_id:
            db.ejecutar_consulta("UPDATE eps SET nombreEPS = %s WHERE id = %s", (texto, eps_id))
        elif accion_eps.get() == "Eliminar" and eps_id:
            db.ejecutar_consulta("DELETE FROM eps WHERE id = %s", (eps_id,))
        
        llenar_tabla_eps()

    marco = ctk.CTkFrame(frame, fg_color="transparent")
    marco.pack(fill="both", expand=True, padx=20, pady=20)

    marco_arl = ctk.CTkFrame(marco, fg_color="transparent")
    marco_arl.pack(side="left", fill="both", expand=True, padx=10, pady=10)

    marco_eps = ctk.CTkFrame(marco, fg_color="transparent")
    marco_eps.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    # ARL Inputs
    ctk.CTkLabel(marco_arl, text="Nombre ARL:", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_arl, textvariable=nombre_arl, text_color=COLOR_TEXTO, fg_color="white").grid(row=0, column=1, padx=5, pady=5)
    ctk.CTkLabel(marco_arl, text="Acción:", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5)
    ctk.CTkComboBox(marco_arl, variable=ejecutar_accion_arl, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=1, column=1, padx=5, pady=5)
    ctk.CTkButton(marco_arl, text="Ejecutar", command=ejecutar_accion_arl, fg_color=COLOR_BOTON).grid(row=2, column=0, columnspan=2, pady=10)

    # ARL Table + Scrollbar
    frame_arl_tabla = ctk.CTkFrame(marco_arl)
    frame_arl_tabla.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvARL = ttk.Treeview(frame_arl_tabla, columns=("ID", "Nombre"), show="headings", height=15)
    tvARL.heading("Nombre", text="Nombre ARL")
    tvARL.column("ID", width=0, stretch=False)  # Oculta la columna de ID
    tvARL.column("Nombre", width=250)
    tvARL.pack(side="left", fill="both", expand=True)
    tvARL.bind("<<TreeviewSelect>>", seleccionar_arl)

    scrollbar_arl = ttk.Scrollbar(frame_arl_tabla, orient="vertical", command=tvARL.yview)
    tvARL.configure(yscrollcommand=scrollbar_arl.set)
    scrollbar_arl.pack(side="right", fill="y")

    # EPS Inputs
    ctk.CTkLabel(marco_eps, text="Nombre EPS:", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_eps, textvariable=nombre_eps, text_color=COLOR_TEXTO, fg_color="white").grid(row=0, column=1, padx=5, pady=5)
    ctk.CTkLabel(marco_eps, text="Acción:", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5)
    ctk.CTkComboBox(marco_eps, variable=ejecutar_accion_eps, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=1, column=1, padx=5, pady=5)
    ctk.CTkButton(marco_eps, text="Ejecutar", command=ejecutar_accion_eps, fg_color=COLOR_BOTON).grid(row=2, column=0, columnspan=2, pady=10)

    # EPS Table + Scrollbar
    frame_eps_tabla = ctk.CTkFrame(marco_eps)
    frame_eps_tabla.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvEPS = ttk.Treeview(frame_eps_tabla, columns=("ID", "Nombre"), show="headings", height=15)
    tvEPS.heading("Nombre", text="Nombre EPS")
    tvEPS.column("ID", width=0, stretch=False)  # Oculta la columna de ID
    tvEPS.column("Nombre", width=250)
    tvEPS.pack(side="left", fill="both", expand=True)
    tvEPS.bind("<<TreeviewSelect>>", seleccionar_eps)

    scrollbar_eps = ttk.Scrollbar(frame_eps_tabla, orient="vertical", command=tvEPS.yview)
    tvEPS.configure(yscrollcommand=scrollbar_eps.set)
    scrollbar_eps.pack(side="right", fill="y")

    llenar_tabla_arl()
    llenar_tabla_eps()

####################################### Pantalla de banco #####################################################
def mostrar_bancos(frame):
    nombre_banco = ctk.StringVar()
    tipo_cuenta = ctk.StringVar()
    accion = ctk.StringVar(value="Agregar")
    banco_id = None  # Para almacenar el ID del banco seleccionado

    def limpiar():
        nonlocal banco_id
        banco_id = None
        nombre_banco.set("")
        tipo_cuenta.set("")
        accion.set("Agregar")

    def llenar_tabla():
        tvBancos.delete(*tvBancos.get_children())
        try:
            db.cursor.execute("""
                SELECT b.id, b.nombreBanco, tc.nombreTipoCuenta
                FROM banco b
                JOIN tipodecuenta tc ON b.tipoCuenta = tc.id
            """)
            for fila in db.cursor.fetchall():
                tvBancos.insert("", "end", values=(fila[0], fila[1], fila[2]))
        except Exception as e:
            print(f"Error al llenar la tabla: {e}")

    def seleccionar(event):
        nonlocal banco_id
        seleccion = tvBancos.selection()
        if seleccion:
            valores = tvBancos.item(seleccion[0], "values")
            if valores:
                banco_id = valores[0]  # Guardamos el ID del banco seleccionado
                nombre_banco.set(valores[1])
                tipo_cuenta.set(valores[2])
                accion.set("Modificar")

    def obtener_o_crear_tipo_cuenta(nombre_tipo):
        """Verifica si el tipo de cuenta existe; si no, lo crea."""
        db.cursor.execute("SELECT id FROM tipodecuenta WHERE nombreTipoCuenta = %s", (nombre_tipo,))
        resultado = db.cursor.fetchone()
        if resultado:
            return resultado[0]
        else:
            db.cursor.execute("INSERT INTO tipodecuenta (nombreTipoCuenta) VALUES (%s)", (nombre_tipo,))
            db.connection.commit()
            return db.cursor.lastrowid  # Retorna el nuevo ID insertado

    def ejecutar_accion():
        nonlocal banco_id
        opcion = accion.get().strip()
        
        if not nombre_banco.get() or not tipo_cuenta.get():
            return

        # Obtener o crear el tipo de cuenta
        id_tipo_cuenta = obtener_o_crear_tipo_cuenta(tipo_cuenta.get())

        if opcion == "Agregar":
            db.ejecutar_consulta("INSERT INTO banco (nombreBanco, tipoCuenta) VALUES (%s, %s)", (nombre_banco.get(), id_tipo_cuenta))
        
        elif opcion == "Modificar" and banco_id:
            db.ejecutar_consulta("UPDATE banco SET nombreBanco=%s, tipoCuenta=%s WHERE id=%s", (nombre_banco.get(), id_tipo_cuenta, banco_id))
        
        elif opcion == "Eliminar" and banco_id:
            db.ejecutar_consulta("DELETE FROM banco WHERE id=%s", (banco_id,))

        llenar_tabla()
        limpiar()

    marco_bancos = ctk.CTkFrame(frame, fg_color="transparent")
    marco_bancos.pack(fill="both", expand=True, padx=20, pady=20)

    # Campo de entrada para el nombre del banco
    ctk.CTkLabel(marco_bancos, text="Nombre del Banco", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5, sticky="w")
    ctk.CTkEntry(marco_bancos, textvariable=nombre_banco, fg_color="white", text_color=COLOR_TEXTO).grid(row=0, column=1, padx=5, pady=5, sticky="ew")

    # Campo de entrada para el tipo de cuenta (ya no es un combobox)
    ctk.CTkLabel(marco_bancos, text="Tipo de Cuenta", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5, sticky="w")
    ctk.CTkEntry(marco_bancos, textvariable=tipo_cuenta, fg_color="white", text_color=COLOR_TEXTO).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

    # Campo de selección para la acción
    ctk.CTkLabel(marco_bancos, text="Acción", text_color=COLOR_TEXTO).grid(row=2, column=0, padx=5, pady=5, sticky="w")
    ctk.CTkComboBox(marco_bancos, variable=ejecutar_accion, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=2, column=1, padx=5, pady=5, sticky="ew")

    # Botón de ejecución
    ctk.CTkButton(marco_bancos, text="Ejecutar", command=ejecutar_accion, fg_color=COLOR_BOTON).grid(row=3, column=0, columnspan=2, pady=10, sticky="ew")

    # Tabla (Treeview) con "Nombre" y "Tipo de Cuenta"
    frame_tabla = ctk.CTkFrame(marco_bancos)
    frame_tabla.grid(row=4, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvBancos = ctk.CTkTabview(frame_tabla, columns=("ID", "Nombre", "Tipo de Cuenta"), show="headings", height=10)
    tvBancos.heading("Nombre", text="Nombre")
    tvBancos.heading("Tipo de Cuenta", text="Tipo de Cuenta")

    # Ocultar la columna de ID
    tvBancos.column("ID", width=0, stretch=False)
    tvBancos.column("Nombre", width=200)
    tvBancos.column("Tipo de Cuenta", width=150)

    tvBancos.pack(side="left", fill="both", expand=True)
    tvBancos.bind("<<TreeviewSelect>>", seleccionar)

    scrollbar = ttk.Scrollbar(frame_tabla, orient="vertical", command=tvBancos.yview)
    tvBancos.configure(yscrollcommand=scrollbar.set)
    scrollbar.pack(side="right", fill="y")

    llenar_tabla()

############################################# Pantalla Ciudad ####################################################
def mostrar_ciudades_departamentos(frame, db):
    nombre_ciudad = ctk.StringVar()
    accion_ciudad = ctk.StringVar(value="Agregar")
    ciudad_id = None

    nombre_departamento = ctk.StringVar()
    accion_departamento = ctk.StringVar(value="Agregar")
    departamento_id = None

    def seleccionar_ciudad(event):
        nonlocal ciudad_id
        seleccion = tvCiudad.selection()
        if seleccion:
            valores = tvCiudad.item(seleccion[0], "values")
            if valores:
                ciudad_id = valores[0]
                nombre_ciudad.set(valores[1])

    def seleccionar_departamento(event):
        nonlocal departamento_id
        seleccion = tvDepartamento.selection()
        if seleccion:
            valores = tvDepartamento.item(seleccion[0], "values")
            if valores:
                departamento_id = valores[0]
                nombre_departamento.set(valores[1])

    def llenar_tabla_ciudad():
        tvCiudad.delete(*tvCiudad.get_children())
        for fila in db.obtener_datos("SELECT id, nombreCiudad FROM ciudad"):
            tvCiudad.insert("", "end", values=fila)

    def llenar_tabla_departamento():
        tvDepartamento.delete(*tvDepartamento.get_children())
        for fila in db.obtener_datos("SELECT id, nombreDepartamento FROM departamentos"):
            tvDepartamento.insert("", "end", values=fila)

    def ejecutar_accion_ciudad():
        if accion_ciudad.get() == "Agregar" and nombre_ciudad.get().strip():
            db.ejecutar_consulta("INSERT INTO ciudad (nombreCiudad) VALUES (%s)", (nombre_ciudad.get(),))
        elif accion_ciudad.get() == "Modificar" and ciudad_id:
            db.ejecutar_consulta("UPDATE ciudad SET nombreCiudad=%s WHERE id=%s", (nombre_ciudad.get(), ciudad_id))
        elif accion_ciudad.get() == "Eliminar" and ciudad_id:
            db.ejecutar_consulta("DELETE FROM ciudad WHERE id=%s", (ciudad_id,))
        llenar_tabla_ciudad()

    def ejecutar_accion_departamento():
        if accion_departamento.get() == "Agregar" and nombre_departamento.get().strip():
            db.ejecutar_consulta("INSERT INTO departamentos (nombreDepartamento) VALUES (%s)", (nombre_departamento.get(),))
        elif accion_departamento.get() == "Modificar" and departamento_id:
            db.ejecutar_consulta("UPDATE departamentos SET nombreDepartamento=%s WHERE id=%s", (nombre_departamento.get(), departamento_id))
        elif accion_departamento.get() == "Eliminar" and departamento_id:
            db.ejecutar_consulta("DELETE FROM departamentos WHERE id=%s", (departamento_id,))
        llenar_tabla_departamento()

    marco = ctk.CTkFrame(frame, fg_color="transparent")
    marco.pack(fill="both", expand=True, padx=20, pady=20)

    marco_ciudad = ctk.CTkFrame(marco, fg_color="transparent")
    marco_ciudad.pack(side="left", fill="both", expand=True, padx=10, pady=10)

    marco_departamento = ctk.CTkFrame(marco, fg_color="transparent")
    marco_departamento.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    # Ciudad Inputs
    ctk.CTkLabel(marco_ciudad, text="Nombre Ciudad:", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_ciudad, textvariable=nombre_ciudad, text_color=COLOR_TEXTO, fg_color="white").grid(row=0, column=1, padx=5, pady=5)
    ctk.CTkLabel(marco_ciudad, text="Acción:", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5)
    ctk.CTkComboBox(marco_ciudad, variable=ejecutar_accion_ciudad, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=1, column=1, padx=5, pady=5)
    ctk.CTkButton(marco_ciudad, text="Ejecutar", command=ejecutar_accion_ciudad, fg_color=COLOR_BOTON).grid(row=2, column=0, columnspan=2, pady=10)

    # Ciudad Table + Scrollbar
    frame_ciudad_tabla = ctk.CTkFrame(marco_ciudad)
    frame_ciudad_tabla.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvCiudad = ttk.Treeview(frame_ciudad_tabla, columns=("ID", "Nombre"), show="headings", height=15)
    tvCiudad.heading("Nombre", text="Nombre Ciudad")
    tvCiudad.column("ID", width=0, stretch=False)
    tvCiudad.column("Nombre", width=250)
    tvCiudad.pack(side="left", fill="both", expand=True)
    tvCiudad.bind("<<TreeviewSelect>>", seleccionar_ciudad)

    scrollbar_ciudad = ttk.Scrollbar(frame_ciudad_tabla, orient="vertical", command=tvCiudad.yview)
    tvCiudad.configure(yscrollcommand=scrollbar_ciudad.set)
    scrollbar_ciudad.pack(side="right", fill="y")

    # Departamento Inputs
    ctk.CTkLabel(marco_departamento, text="Nombre Departamento:", text_color="black").grid(row=0, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_departamento, textvariable=nombre_departamento, text_color="black", fg_color="white").grid(row=0, column=1, padx=5, pady=5)
    ctk.CTkLabel(marco_departamento, text="Acción:", text_color="black").grid(row=1, column=0, padx=5, pady=5)
    ctk.CTkComboBox(marco_departamento, variable=ejecutar_accion_departamento, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=1, column=1, padx=5, pady=5)
    ctk.CTkButton(marco_departamento, text="Ejecutar", command=ejecutar_accion_departamento, fg_color=COLOR_BOTON).grid(row=2, column=0, columnspan=2, pady=10)

    # Departamento Table + Scrollbar
    frame_departamento_tabla = ctk.CTkFrame(marco_departamento)
    frame_departamento_tabla.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvDepartamento = ttk.Treeview(frame_departamento_tabla, columns=("ID", "Nombre"), show="headings", height=15)
    tvDepartamento.heading("Nombre", text="Nombre Departamento")
    tvDepartamento.column("ID", width=0, stretch=False)
    tvDepartamento.column("Nombre", width=250)
    tvDepartamento.pack(side="left", fill="both", expand=True)
    tvDepartamento.bind("<<TreeviewSelect>>", seleccionar_departamento)

    scrollbar_departamento = ttk.Scrollbar(frame_departamento_tabla, orient="vertical", command=tvDepartamento.yview)
    tvDepartamento.configure(yscrollcommand=scrollbar_departamento.set)
    scrollbar_departamento.pack(side="right", fill="y")

    llenar_tabla_ciudad()
    llenar_tabla_departamento()

########################################### Pantalla Cargo #####################################################
def mostrar_cargo(frame,db):
    global records_per_page, page_number
    
    db = DataBase()
    records_per_page = 10
    page_number = 1

    # Variables de control
    nombre_cargo = ctk.StringVar()
    nombre_jefe = ctk.StringVar()
    accion = ctk.StringVar(value="Agregar")  # Opción por defecto

    def seleccionar(event):
        seleccion = tvCargos.selection()
        if seleccion:
            valores = tvCargos.item(seleccion[0], "values")
            nombre_cargo.set(valores[0])
            nombre_jefe.set(valores[1])

    def limpiar():
        nombre_cargo.set("")
        nombre_jefe.set("")

    def vaciar_tabla():
        for fila in tvCargos.get_children():
            tvCargos.delete(fila)

    def llenar_tabla():
        vaciar_tabla()
        sql = "SELECT c.nombreCargo, j.nombreJefe FROM cargo c JOIN jefe j ON c.idJefe = j.id"
        db.cursor.execute(sql)
        for fila in db.cursor.fetchall():
            tvCargos.insert("", "end", values=(fila[0], fila[1]))

    def ejecutar_accion():
        seleccionada = accion.get()

        if seleccionada == "Agregar":
            if nombre_cargo.get() and nombre_jefe.get():
                # Insertar jefe si no existe
                db.cursor.execute(
                    "INSERT INTO jefe (nombreJefe) VALUES (%s) ON DUPLICATE KEY UPDATE nombreJefe=VALUES(nombreJefe)",
                    (nombre_jefe.get(),)
                )
                db.connection.commit()

                # Obtener el ID del jefe
                db.cursor.execute("SELECT id FROM jefe WHERE nombreJefe=%s", (nombre_jefe.get(),))
                resultado = db.cursor.fetchone()
                if resultado:
                    id_jefe = resultado[0]
                else:
                    print("⚠️ Error inesperado: No se encontró el jefe en la base de datos.")
                    return  # Salir sin continuar

                # Insertar cargo
                db.cursor.execute("INSERT INTO cargo (nombreCargo, idJefe) VALUES (%s, %s)",
                                (nombre_cargo.get(), id_jefe))
                db.connection.commit()
                llenar_tabla()
                limpiar()

        elif seleccionada == "Modificar":
            seleccion = tvCargos.selection()
            if seleccion and nombre_cargo.get() and nombre_jefe.get():
                valores = tvCargos.item(seleccion[0], "values")

                # Intentar obtener el ID del jefe
                db.cursor.execute("SELECT id FROM jefe WHERE nombreJefe=%s", (nombre_jefe.get(),))
                resultado = db.cursor.fetchone()

                if resultado:
                    id_jefe = resultado[0]  # Si existe, obtenemos su ID
                else:
                    # Si no existe, lo insertamos en la base de datos
                    db.cursor.execute("INSERT INTO jefe (nombreJefe) VALUES (%s)", (nombre_jefe.get(),))
                    db.connection.commit()

                    # Obtener el nuevo ID del jefe
                    db.cursor.execute("SELECT id FROM jefe WHERE nombreJefe=%s", (nombre_jefe.get(),))
                    resultado = db.cursor.fetchone()
                    if resultado:
                        id_jefe = resultado[0]
                    else:
                        print("⚠️ Error inesperado: No se pudo crear el jefe.")
                        return

                # Actualizar el cargo con el nuevo ID del jefe
                db.cursor.execute("UPDATE cargo SET nombreCargo=%s, idJefe=%s WHERE nombreCargo=%s",
                                (nombre_cargo.get(), id_jefe, valores[0]))
                db.connection.commit()
                llenar_tabla()
                limpiar()

        elif seleccionada == "Eliminar":
            seleccion = tvCargos.selection()
            if seleccion:
                valores = tvCargos.item(seleccion[0], "values")
                db.cursor.execute("DELETE FROM cargo WHERE nombreCargo=%s", (valores[0],))
                db.connection.commit()
                llenar_tabla()
                limpiar()

    marco_cargo = ctk.CTkFrame(frame, fg_color="transparent")
    marco_cargo.pack(fill="both", expand=True, padx=20, pady=20)

    # Contenedor central
    contenedor = ctk.CTkFrame(marco_cargo, fg_color="transparent")
    contenedor.pack(expand=True, pady=20)

    # Campos de entrada
    ctk.CTkLabel(contenedor, text="Nombre Cargo", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=10, pady=5, sticky="w")
    txtNombreCargo = ctk.CTkEntry(contenedor, textvariable=nombre_cargo, fg_color="white", text_color=COLOR_TEXTO)
    txtNombreCargo.grid(row=0, column=1, padx=10, pady=5, sticky="ew")

    ctk.CTkLabel(contenedor, text="Jefe", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=10, pady=5, sticky="w")
    txtJefe = ctk.CTkEntry(contenedor, textvariable=nombre_jefe, fg_color="white", text_color=COLOR_TEXTO)
    txtJefe.grid(row=1, column=1, padx=10, pady=5, sticky="ew")

    # Selección de acción
    ctk.CTkLabel(contenedor, text="Acción", text_color=COLOR_TEXTO).grid(row=2, column=0, padx=10, pady=5, sticky="w")
    comboAccion = ctk.CTkComboBox(contenedor, variable=ejecutar_accion, fg_color="white", values=["Agregar", "Modificar", "Eliminar"])
    comboAccion.grid(row=2, column=1, padx=10, pady=5, sticky="ew")

    # Botón de ejecutar
    btnEjecutar = ctk.CTkButton(contenedor, text="Ejecutar", command=ejecutar_accion, fg_color=COLOR_BOTON)
    btnEjecutar.grid(row=3, column=0, columnspan=2, pady=10, sticky="ew")

    # Tabla
    tvCargos = ttk.Treeview(marco_cargo, columns=("Nombre Cargo", "Jefe"), show="headings")
    tvCargos.heading("Nombre Cargo", text="Nombre Cargo")
    tvCargos.heading("Jefe", text="Jefe")
    tvCargos.pack(fill="both", expand=True, padx=10, pady=10)
    tvCargos.bind("<<TreeviewSelect>>", seleccionar)

    llenar_tabla()

##################################### Tipo de contrato ##################################################################
def mostrar_tipodecontrato(frame):
    global db
    db = DataBase()
    tipo_contrato = ctk.StringVar()
    accion = ctk.StringVar(value="Agregar")
    contrato_id = None  # Para almacenar el ID del tipo de contrato seleccionado

    def limpiar():
        nonlocal contrato_id
        contrato_id = None
        tipo_contrato.set("")
        accion.set("Agregar")

    def llenar_tabla():
        tvContratos.delete(*tvContratos.get_children())
        try:
            db.cursor.execute("SELECT id, nombreTipoContrato FROM tipodecontrato")
            for fila in db.cursor.fetchall():
                tvContratos.insert("", "end", values=(fila[0], fila[1]))
        except Exception as e:
            print(f"Error al llenar la tabla: {e}")

    def seleccionar(event):
        nonlocal contrato_id
        seleccion = tvContratos.selection()
        if seleccion:
            valores = tvContratos.item(seleccion[0], "values")
            if valores:
                contrato_id = valores[0]
                tipo_contrato.set(valores[1])
                accion.set("Modificar")

    def ejecutar_accion():
        nonlocal contrato_id
        opcion = accion.get().strip()

        if not tipo_contrato.get():
            return

        if opcion == "Agregar":
            db.ejecutar_consulta("INSERT INTO tipodecontrato (nombreTipoContrato) VALUES (%s)", (tipo_contrato.get(),))
        elif opcion == "Modificar" and contrato_id:
            db.ejecutar_consulta("UPDATE tipodecontrato SET nombreTipoContrato=%s WHERE id=%s", (tipo_contrato.get(), contrato_id))
        elif opcion == "Eliminar" and contrato_id:
            db.ejecutar_consulta("DELETE FROM tipodecontrato WHERE id=%s", (contrato_id,))

        llenar_tabla()
        limpiar()

    marco_contrato = ctk.CTkFrame(frame, fg_color="transparent")
    marco_contrato.pack(fill="both", expand=True, padx=20, pady=20)

    ctk.CTkLabel(marco_contrato, text="Tipo de Contrato", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5, sticky="w")
    ctk.CTkEntry(marco_contrato, textvariable=tipo_contrato, fg_color="white", text_color=COLOR_TEXTO).grid(row=0, column=1, padx=5, pady=5, sticky="ew")

    ctk.CTkLabel(marco_contrato, text="Acción", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5, sticky="w")
    ctk.CTkComboBox(marco_contrato, variable=ejecutar_accion, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=1, column=1, padx=5, pady=5, sticky="ew")

    ctk.CTkButton(marco_contrato, text="Ejecutar", command=ejecutar_accion, fg_color=COLOR_BOTON).grid(row=2, column=0, columnspan=2, pady=10, sticky="ew")

    frame_tabla = ctk.CTkFrame(marco_contrato)
    frame_tabla.grid(row=3, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)

    tvContratos = ttk.Treeview(frame_tabla, columns=("ID", "Tipo de Contrato"), show="headings", height=10)
    tvContratos.heading("Tipo de Contrato", text="Tipo de Contrato")
    
    tvContratos.column("ID", width=0, stretch=False)
    tvContratos.column("Tipo de Contrato", width=200)
    
    tvContratos.pack(side="left", fill="both", expand=True)
    tvContratos.bind("<<TreeviewSelect>>", seleccionar)

    scrollbar = ttk.Scrollbar(frame_tabla, orient="vertical", command=tvContratos.yview)
    tvContratos.configure(yscrollcommand=scrollbar.set)
    scrollbar.pack(side="right", fill="y")

    llenar_tabla()


########################################### Pantalla Areas ########################################################3
def mostrar_areas(frame):
    global db
    db = DataBase()
    nombre_dependencia = ctk.StringVar()
    accion = ctk.StringVar(value="Agregar")

    def vaciar_tabla():
        for fila in tvDependencias.get_children():
            tvDependencias.delete(fila)

    def limpiar():
        nombre_dependencia.set("")
        accion.set("Agregar")

    def llenar_tabla():
        vaciar_tabla()
        try:
            sql = "SELECT nombreDependencia FROM dependencia"
            db.cursor.execute(sql)
            for fila in db.cursor.fetchall():
                tvDependencias.insert("", 'end', values=(fila[0],))
        except Exception as e:
            print(f"Error al llenar la tabla: {e}")

    def seleccionar(event):
        seleccion = tvDependencias.selection()
        if seleccion:
            nombre = tvDependencias.item(seleccion[0], "values")[0]
            nombre_dependencia.set(nombre)

    def ejecutar():
        if accion.get() == "Agregar":
            sql = "INSERT INTO dependencia (nombreDependencia) VALUES (%s)"
            db.cursor.execute(sql, (nombre_dependencia.get(),))
        elif accion.get() == "Modificar":
            seleccion = tvDependencias.selection()
            if seleccion:
                nombre_actual = tvDependencias.item(seleccion[0], "values")[0]
                sql = "UPDATE dependencia SET nombreDependencia=%s WHERE nombreDependencia=%s"
                db.cursor.execute(sql, (nombre_dependencia.get(), nombre_actual))
        elif accion.get() == "Eliminar":
            seleccion = tvDependencias.selection()
            if seleccion:
                nombre_actual = tvDependencias.item(seleccion[0], "values")[0]
                sql = "DELETE FROM dependencia WHERE nombreDependencia=%s"
                db.cursor.execute(sql, (nombre_actual,))
                tvDependencias.delete(seleccion[0])
        db.connection.commit()
        llenar_tabla()
        limpiar()

    marco = ctk.CTkFrame(frame, fg_color="transparent")
    marco.pack(fill="both", expand=True, padx=20, pady=20)

    ctk.CTkLabel(marco, text="Nombre", text_color=COLOR_TEXTO).grid(column=0, row=0, padx=5, pady=5)
    ctk.CTkEntry(marco, textvariable=nombre_dependencia, fg_color="white", text_color=COLOR_TEXTO).grid(column=1, row=0)

    ctk.CTkLabel(marco, text="Acción", text_color=COLOR_TEXTO).grid(column=0, row=1, padx=5, pady=5)
    cbAccion = ctk.CTkComboBox(marco, variable=accion, fg_color="white", values=["Agregar", "Modificar", "Eliminar"])
    cbAccion.grid(column=1, row=1)

    ctk.CTkButton(marco, text="Ejecutar", command=ejecutar, fg_color=COLOR_BOTON).grid(column=0, row=3, columnspan=2, pady=10)

    frame_tabla = ctk.CTkFrame(marco)
    frame_tabla.grid(column=0, row=2, columnspan=2, padx=5, pady=5)

    scrollbar = ttk.Scrollbar(frame_tabla, orient="vertical")
    scrollbar.pack(side="right", fill="y")

    tvDependencias = ttk.Treeview(frame_tabla, selectmode='browse', columns=("Nombre",), yscrollcommand=scrollbar.set)
    tvDependencias.column("#0", width=0, stretch='no')
    tvDependencias.column("Nombre", width=300)
    tvDependencias.heading("#0", text="")
    tvDependencias.heading("Nombre", text="Nombre", anchor='center')
    tvDependencias.pack(side="left", fill="both", expand=True)
    tvDependencias.bind("<<TreeviewSelect>>", seleccionar)

    scrollbar.config(command=tvDependencias.yview)

    llenar_tabla()

def formatear_fecha_larga(fecha):
    return fecha.strftime("%d de %B de %Y")  # Formato: 22 de enero de 2024

def mostrar_certificado(frame):
    db = DataBase()
    modificar = False
    descripcion_contrato = tk.StringVar()
    autorizacion_contratos = tk.StringVar()
    valor_contrato = tk.StringVar()
    obligaciones_especificas = tk.StringVar()
    consecutivo = tk.StringVar()
    id_independiente = tk.StringVar()
    id_version = tk.StringVar()
    id_objeto = tk.StringVar()

    # Variables de paginación
    page_number = 1
    records_per_page = 10

    # Variables de filtro
    filtro_cliente = tk.StringVar()
    filtro_tipo_contrato = tk.StringVar()
    filtro_cargo = tk.StringVar()
    filtro_dependencia = tk.StringVar()
    filtro_autorizacion = tk.StringVar()
    filtro_valor = tk.StringVar()
    filtro_ObEspCon = tk.StringVar()
    filtro_consecutivo = tk.StringVar()
    filtro_id_version = tk.StringVar()
    filtro_id_comboObjeto = tk.StringVar()
    filtro_id_Obligaciones = tk.StringVar()

    def cargar_combo_data():
        try:
            # Tipo de Objeto
            sql = "SELECT nombre FROM obligacionesespecificas"
            db.cursor.execute(sql)
            obligaciones_especificas = db.cursor.fetchall()
            comboObligaciones['values'] = [row[0] for row in obligaciones_especificas]
            comboFiltroObligaciones_Específicas['values'] = comboObligaciones['values']  # También cargar en filtro

            # Tipo de Objeto
            sql = "SELECT nombreObjeto FROM objeto"
            db.cursor.execute(sql)
            objeto = db.cursor.fetchall()
            comboObjeto['values'] = [row[0] for row in objeto]
            comboFiltroTipoContrato['values'] = comboObjeto['values']  # También cargar en filtro
            
            # Tipo de Contrato
            sql = "SELECT nombreTipoContrato FROM tipodecontrato"
            db.cursor.execute(sql)
            tipodecontrato = db.cursor.fetchall()
            comboTipoContrato['values'] = [row[0] for row in tipodecontrato]
            comboFiltroTipoContrato['values'] = comboTipoContrato['values']  # También cargar en filtro

            # Cargo
            sql = "SELECT nombreCargo FROM cargo"
            db.cursor.execute(sql)
            cargos = db.cursor.fetchall()
            comboCargo['values'] = [row[0] for row in cargos]
            comboFiltroCargo['values'] = comboCargo['values']  # También cargar en filtro

            # Dependencia
            sql = "SELECT nombreDependencia FROM dependencia"
            db.cursor.execute(sql)
            dependencias = db.cursor.fetchall()
            comboDependencia['values'] = [row[0] for row in dependencias]
            comboFiltroDependencia['values'] = comboDependencia['values'] 

            # Clientes - Solo mostrar clientes activos
            sql = """
            SELECT clientes.id AS cliente_id, CONCAT_WS(' ', primerNombre, segundoNombre, primerApellido, segundoApellido) AS nombreCompleto 
            FROM clientes
            JOIN estado ON clientes.estado = estado.id
            WHERE estado.tipoEstado = 'activo'

            """
            db.cursor.execute(sql)
            clientes = db.cursor.fetchall()
            comboClientes['values'] = [row[1] for row in clientes]
            comboFiltroCliente['values'] = comboClientes['values']

        except Exception as e:
            print(f"Error al cargar datos en combobox: {e}")

    def limpiar_campos():
        id_independiente.set("")
        descripcion_contrato.set("")
        autorizacion_contratos.set("")
        valor_contrato.set("")
        comboObjeto.set("")
        obligaciones_especificas.set("")
        comboClientes.set("")
        comboTipoContrato.set("")
        comboCargo.set("")
        comboDependencia.set("")
        date_vigencia.set_date(datetime.today())
        date_terminacion.set_date(datetime.today())
        consecutivo.set("")
        comboObligaciones.set("")

    def llenar_tabla():
        print("Cargando datos en la tabla...")
        for item in tabla.get_children():
            tabla.delete(item)

        try:
            offset = (page_number - 1) * records_per_page
            # Base SQL query
            sql = """
            SELECT contrato.id, contrato.idVersion, clientes.id,
                CONCAT_WS(' ', clientes.primerNombre, clientes.segundoNombre, clientes.primerApellido, clientes.segundoApellido) AS nombreCompleto,
                tipodecontrato.nombreTipoContrato, cargo.nombreCargo, dependencia.nombreDependencia, objeto.nombreObjeto,
                contrato.descripcionContrato, contrato.Vigencia, contrato.terminacion, contrato.autorizacionContratos, contrato.valorContrato,
                obligacionesespecificas.nombre AS nombreObligacion, contrato.consecutivo
            FROM contrato
            JOIN clientes ON contrato.idClientes = clientes.id
            JOIN tipodecontrato ON contrato.idTipoContrato = tipodecontrato.id
            JOIN cargo ON contrato.idCargo = cargo.id
            JOIN dependencia ON contrato.idDependecia = dependencia.id
            JOIN objeto ON contrato.idObjeto = objeto.id
            JOIN obligacionesespecificas ON contrato.ObEspCon = obligacionesespecificas.id
            WHERE 1=1
            """

            
            valores = []

            # Aplicar filtros si se ingresaron valores
            if filtro_id_version.get():
                sql += " AND contrato.idVersion LIKE %s"
                valores.append(f"%{filtro_id_version.get()}%")
            if filtro_cliente.get():
                sql += " AND CONCAT_WS(' ', clientes.primerNombre, clientes.segundoNombre, clientes.primerApellido, clientes.segundoApellido) LIKE %s"
                valores.append(f"%{filtro_cliente.get()}%")
            if filtro_tipo_contrato.get():
                sql += " AND tipodecontrato.nombreTipoContrato LIKE %s"
                valores.append(f"%{filtro_tipo_contrato.get()}%")
            if filtro_cargo.get():
                sql += " AND cargo.nombreCargo LIKE %s"
                valores.append(f"%{filtro_cargo.get()}%")
            if filtro_dependencia.get():
                sql += " AND dependencia.nombreDependencia LIKE %s"
                valores.append(f"%{filtro_dependencia.get()}%")
                
            if filtro_autorizacion.get():
                sql += " AND contrato.autorizacionContratos LIKE %s"
                valores.append(f"%{filtro_autorizacion.get()}%")
            if filtro_valor.get():
                sql += " AND contrato.valorContrato LIKE %s"
                valores.append(f"%{filtro_valor.get()}%")

            # Agregar la paginación a la consulta
            sql += " LIMIT %s OFFSET %s"
            valores.extend([records_per_page, offset])

            db.cursor.execute(sql, valores)
            contratos = db.cursor.fetchall()
            print(f"Contratos obtenidos: {contratos}")

            for contrato in contratos:
                tabla.insert('', 'end', values=contrato)

        except Exception as e:  
            print(f"Error al llenar la tabla: {e}")

    def seleccionar_contrato(event):
        nonlocal modificar
        modificar = True

        selected_item = tabla.selection()
        if selected_item:
            item = tabla.item(selected_item)
            contrato = item['values']

            # Asignar valores a las variables usando .set() para cada StringVar
            id_independiente.set(contrato[0])
            id_version.set(contrato[1])
            comboClientes.set(contrato[3])
            comboTipoContrato.set(contrato[4])
            comboCargo.set(contrato[5])
            comboDependencia.set(contrato[6])
            comboObjeto.set(contrato[7])
            descripcion_contrato.set(contrato[8])
            date_vigencia.set_date(contrato[9])
            date_terminacion.set_date(contrato[10])
            autorizacion_contratos.set(contrato[11])
            valor_contrato.set(contrato[12])
            comboObligaciones.set(contrato[13])
            consecutivo.set(contrato[14])


    def guardar_contrato():
        if modificar:
            actualizar_contrato()
        else:
            insertar_contrato()

    def insertar_contrato():
        try:
            if not date_vigencia.get_date() or not date_terminacion.get_date():
                print("Error: Las fechas de vigencia y terminación son obligatorias.")
                return

            # Obtener los IDs necesarios para la inserción
            sql_cliente = """
            SELECT id FROM clientes 
            WHERE CONCAT_WS(' ', primerNombre, segundoNombre, primerApellido, segundoApellido) = %s
            LIMIT 1
            """
            db.cursor.execute(sql_cliente, (comboClientes.get(),))
            cliente_id = db.cursor.fetchone()

            sql_tipo_contrato = """
            SELECT id FROM tipodecontrato 
            WHERE nombreTipoContrato = %s
            LIMIT 1
            """
            db.cursor.execute(sql_tipo_contrato, (comboTipoContrato.get(),))
            tipo_contrato_id = db.cursor.fetchone()

            sql_cargo = """
            SELECT id FROM cargo 
            WHERE nombreCargo = %s
            LIMIT 1
            """
            db.cursor.execute(sql_cargo, (comboCargo.get(),))
            cargo_id = db.cursor.fetchone()

            sql_dependencia = """
            SELECT id FROM dependencia 
            WHERE nombreDependencia = %s
            LIMIT 1
            """
            db.cursor.execute(sql_dependencia, (comboDependencia.get(),))
            dependencia_id = db.cursor.fetchone()

            sql_objeto = """
            SELECT id FROM objeto    
            WHERE nombreObjeto = %s
            LIMIT 1
            """
            db.cursor.execute(sql_objeto, (comboObjeto.get(),))
            objeto_id = db.cursor.fetchone()

            sql_ObEspCon = """
            SELECT id FROM obligacionesespecificas    
            WHERE nombre = %s
            LIMIT 1
            """
            db.cursor.execute(sql_ObEspCon, (comboObligaciones.get(),))
            obligaciones_especificas = db.cursor.fetchone()

            # Validar que todos los IDs se hayan encontrado
            if cliente_id and tipo_contrato_id and cargo_id and dependencia_id and objeto_id and obligaciones_especificas:
                # Insertar el contrato sin `idVersion`
                sql_insert = """
                INSERT INTO contrato 
                (idClientes, idTipoContrato, idCargo, idDependecia, idObjeto, descripcionContrato, Vigencia, terminacion, autorizacionContratos, valorContrato, ObEspCon, consecutivo) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                valores = (
                    cliente_id[0],
                    tipo_contrato_id[0],
                    cargo_id[0],
                    dependencia_id[0],
                    objeto_id[0],
                    descripcion_contrato.get(),
                    date_vigencia.get_date(),
                    date_terminacion.get_date(),
                    autorizacion_contratos.get(),
                    valor_contrato.get(),
                    obligaciones_especificas[0],  # Acceder al ID desde la tupla
                    consecutivo.get()  # Consecutivo
                )
                db.cursor.execute(sql_insert, valores)
                db.cursor.connection.commit()

                # Obtener el ID del registro recién insertado
                contrato_id = db.cursor.lastrowid

                # Crear el valor para `idVersion` como `id-1`
                id_version = f"{contrato_id}-1"

                # Actualizar el registro con `idVersion`
                sql_update_version = """
                UPDATE contrato
                SET idVersion = %s
                WHERE id = %s
                """
                db.cursor.execute(sql_update_version, (id_version, contrato_id))
                db.cursor.connection.commit()

                limpiar_campos()
                llenar_tabla()
            else:
                print("Error: No se encontraron todos los IDs requeridos.")

        except Exception as e:
            print(f"Error al insertar contrato: {e}")
            db.cursor.connection.rollback()

    def actualizar_contrato():
        try:
            selected_item = tabla.selection()
            if selected_item:
                item = tabla.item(selected_item)
                contrato_id = item['values'][0]

                # Obtener los IDs necesarios para la actualización
                sql_cliente = """
                SELECT id FROM clientes 
                WHERE CONCAT_WS(' ', primerNombre, segundoNombre, primerApellido, segundoApellido) = %s
                LIMIT 1
                """
                db.cursor.execute(sql_cliente, (comboClientes.get(),))
                cliente_id = db.cursor.fetchone()

                sql_tipo_contrato = """
                SELECT id FROM tipodecontrato 
                WHERE nombreTipoContrato = %s
                LIMIT 1
                """
                db.cursor.execute(sql_tipo_contrato, (comboTipoContrato.get(),))
                tipo_contrato_id = db.cursor.fetchone()

                sql_cargo = """
                SELECT id FROM cargo 
                WHERE nombreCargo = %s
                LIMIT 1
                """
                db.cursor.execute(sql_cargo, (comboCargo.get(),))
                cargo_id = db.cursor.fetchone()

                sql_dependencia = """
                SELECT id FROM dependencia 
                WHERE nombreDependencia = %s
                LIMIT 1
                """
                db.cursor.execute(sql_dependencia, (comboDependencia.get(),))
                dependencia_id = db.cursor.fetchone()

                sql_obligaciones_especificas = """
                SELECT id FROM obligacionesespecificas
                WHERE nombre = %s
                LIMIT 1
                """
                db.cursor.execute(sql_obligaciones_especificas, (comboObligaciones.get(),))
                obligaciones_especificas_id = db.cursor.fetchone()

                # Imprimir el valor del comboObjeto
                objeto_nombre = comboObjeto.get()
                print(f"Valor de objeto: {objeto_nombre}")

                sql_objeto = """
                SELECT id FROM objeto
                WHERE nombreObjeto = %s
                LIMIT 1
                """
                db.cursor.execute(sql_objeto, (objeto_nombre,))
                objeto_id = db.cursor.fetchone()

                # Validar que todos los IDs se hayan encontrado
                if cliente_id and tipo_contrato_id and cargo_id and dependencia_id and obligaciones_especificas_id and objeto_id:
                    sql_update = """
                    UPDATE contrato 
                    SET idClientes = %s,
                        idTipoContrato = %s,
                        idCargo = %s,
                        idDependecia = %s,
                        idObjeto = %s,
                        descripcionContrato = %s, 
                        Vigencia = %s,
                        terminacion = %s,
                        autorizacionContratos = %s, 
                        valorContrato = %s,
                        ObEspCon = %s,
                        consecutivo = %s
                    WHERE id = %s
                    """
                    valores = (
                        cliente_id[0],
                        tipo_contrato_id[0],
                        cargo_id[0],
                        dependencia_id[0],
                        objeto_id[0],
                        descripcion_contrato.get(),
                        date_vigencia.get_date(),
                        date_terminacion.get_date(),
                        autorizacion_contratos.get(),
                        valor_contrato.get(),
                        obligaciones_especificas_id[0],
                        consecutivo.get(),
                        contrato_id
                    )
                    db.cursor.execute(sql_update, valores)
                    db.cursor.connection.commit()

                    limpiar_campos()
                    llenar_tabla()
                else:
                    print("Error: No se encontraron todos los IDs requeridos.")
                    if not cliente_id:
                        print(f"No se encontró un cliente con el nombre: {comboClientes.get()}")
                    if not tipo_contrato_id:
                        print(f"No se encontró un tipo de contrato con el nombre: {comboTipoContrato.get()}")
                    if not cargo_id:
                        print(f"No se encontró un cargo con el nombre: {comboCargo.get()}")
                    if not dependencia_id:
                        print(f"No se encontró una dependencia con el nombre: {comboDependencia.get()}")
                    if not obligaciones_especificas_id:
                        print(f"No se encontró una obligación específica con el nombre: {comboObligaciones.get()}")
                    if not objeto_id:
                        print(f"No se encontró un objeto con el nombre: {objeto_nombre}")

        except Exception as e:
            print(f"Error al actualizar contrato: {e}")
            db.cursor.connection.rollback()

    def eliminar_contrato():
        try:
            selected_item = tabla.selection()
            if selected_item:
                item = tabla.item(selected_item)
                contrato_id = item['values'][0]

                # Obtener un bloqueo sobre el registro antes de eliminar
                db.cursor.execute("SELECT * FROM contrato WHERE id = %s FOR UPDATE", (contrato_id,))

                # Eliminar el contrato
                sql = "DELETE FROM contrato WHERE id = %s"
                db.cursor.execute(sql, (contrato_id,))
                db.cursor.connection.commit()

                limpiar_campos()
                llenar_tabla()
                print("Contrato eliminado exitosamente")

        except Exception as e:
            print(f"Error al eliminar contrato: {e}")
            db.cursor.connection.rollback()

    def ir_a_pagina_anterior():
        nonlocal page_number
        if page_number > 1:
            page_number -= 1
            llenar_tabla()
            actualizar_botones_paginacion()

    def ir_a_pagina_siguiente():
        nonlocal page_number
        page_number += 1
        llenar_tabla()
        actualizar_botones_paginacion()

    def mostrar_todos_los_registros():
        filtro_cliente.set("")
        filtro_tipo_contrato.set("")
        filtro_cargo.set("")
        filtro_dependencia.set("")
        filtro_id_comboObjeto.set("")
        filtro_autorizacion.set("")
        filtro_valor.set("")
        filtro_ObEspCon.set("")
        filtro_consecutivo.set("")
        filtro_id_version.set("")
        llenar_tabla()

    def actualizar_botones_paginacion():
        btnAnterior.config(state=tk.NORMAL if page_number > 1 else tk.DISABLED)

    # Función para abrir una nueva pantalla que permita ver la barra de Windows
    def abrir_nueva_pantalla():
        nueva_pantalla = tk.Toplevel()
        nueva_pantalla.title("Nueva pantalla")
        
        # Adaptar la nueva pantalla al tamaño de la pantalla, pero permitiendo ver la barra de tareas
        screen_width = frame.winfo_screenwidth()
        screen_height = frame.winfo_screenheight() - 50  # Dejar espacio para la barra de Windows
        nueva_pantalla.geometry(f"{screen_width}x{screen_height}+0+0")  # Colocar en la esquina superior izquierda

        # Crear un marco para los botones minimizar y cerrar
        botones_marco = tk.Frame(nueva_pantalla)
        botones_marco.pack(pady=10)

        # Botón minimizar (amarillo)
        minimizar_btn = tk.Button(botones_marco, text="—", command=nueva_pantalla.iconify, bg="yellow", fg="black", font=("Helvetica", 14), width=5)
        minimizar_btn.grid(row=0, column=0, padx=5)

        # Botón cerrar (rojo)
        cerrar_btn = tk.Button(botones_marco, text="X", command=nueva_pantalla.destroy, bg="red", fg="white", font=("Helvetica", 14), width=5)
        cerrar_btn.grid(row=0, column=1, padx=5)
    
    # Crear el marco dentro de la pantalla principal
    marco = tk.Frame(frame)
    marco.pack(fill='both', expand=True)

    # Definir la fuente para el título h1
    font_h1 = font.Font(family="Helvetica", size=20, weight="bold")

    # Crear la etiqueta que simula un título h1
    label_h1 = tk.Label(marco, text="Contrataciones", font=font_h1)
    label_h1.pack(pady=3)

    # Crear un marco para organizar los botones debajo del título
    marco_botones = tk.Frame(marco)
    marco_botones.pack(pady=3)
    # Contenedor para los formularios y la tabla
    contenedor_formulario = tk.Frame(marco)
    contenedor_formulario.pack(fill='x', padx=10, pady=10)

    contenedor_fila_id = tk.Frame(contenedor_formulario)
    contenedor_fila_id.pack(fill='x', pady=5)
    # Validación para asegurarse de que el input sea numérico
    def solo_numeros(texto):
        return texto.isdigit()

    validacion_numerica = frame.register(solo_numeros)
    db = DataBase()
    # Crear la conexión a la base de datos y obtener nombres de objetos
    try:
        db = DataBase()
        nombres_objetos = db.obtener_nombres_objetos()
        print("Contratos obtenidos:", nombres_objetos)
    except Exception as e:
        print(f"Error al obtener nombres de objetos: {e}")
        nombres_objetos = []  # Asigna una lista vacía en caso de error
    
    db = DataBase()
    # Crear la conexión a la base de datos y obtener nombres de objetos
    try:
        db = DataBase()
        nombres_obligaciones_especificas= db.obtener_nombres_obligaciones_especificas()
        print("Contratos obtenidos:", nombres_obligaciones_especificas)
    except Exception as e:
        print(f"Error al obtener nombres de obligaciones especificas: {e}")
        nombres_obligaciones_especificas = []  # Asigna una lista vacía en caso de error

    def abrir_objeto_window():
        # Crear una nueva pantalla para "Objeto"
        pantalla_objeto = tk.Toplevel()
        pantalla_objeto.title("CRUD Objeto")

        # Inicializar variables
        db = DataBase()
        global page_number, records_per_page, total_records
        page_number = 1
        records_per_page = 10
        total_records = 0

        id_objeto = tk.StringVar()
        nombre_objeto = tk.StringVar()
        filtro_nombre = tk.StringVar()

        def aplicar_filtro(*args):
            llenar_tabla()

        def seleccionar(event):
            # Actualizar el campo de texto con el registro seleccionado
            seleccion = tvObjetos.selection()
            if seleccion:
                id = seleccion[0]
                valores = tvObjetos.item(id, "values")
                id_objeto.set(valores[0])
                nombre_objeto.set(valores[1])

        def eliminar():
            seleccion = tvObjetos.selection()
            if seleccion:
                id = seleccion[0]
                if int(id) > 0:
                    sql = "DELETE FROM objeto WHERE id=%s"
                    try:
                        db.cursor.execute(sql, (id,))
                        db.connection.commit()
                        tvObjetos.delete(id)
                        lblMensaje.config(text="Se ha eliminado el registro correctamente", fg="green")
                        limpiar()
                        llenar_tabla()  # Actualizar tabla después de eliminar
                    except pymysql.MySQLError as e:
                        lblMensaje.config(text=f"Error al eliminar el registro: {e}", fg="red")
                else:
                    lblMensaje.config(text="Seleccione un registro para eliminar", fg="red")

        def nuevo():
            if validar():
                if id_objeto.get() == "":  # Si no hay ID, es un nuevo registro
                    # Obtener el siguiente ID auto-incremental
                    sql_get_id = "SELECT IFNULL(MAX(id), 0) + 1 FROM objeto"
                    db.cursor.execute(sql_get_id)
                    nuevo_id = db.cursor.fetchone()[0]

                    sql = "INSERT INTO objeto (id, nombreObjeto) VALUES (%s, %s)"
                    try:
                        db.cursor.execute(sql, (nuevo_id, nombre_objeto.get()))
                        db.connection.commit()
                        lblMensaje.config(text="Se ha guardado el registro con éxito", fg="green")
                        llenar_tabla()
                        limpiar()
                    except pymysql.MySQLError as e:
                        lblMensaje.config(text=f"Error al guardar el registro: {e}", fg="red")
                else:  # Si hay ID, es una actualización
                    actualizar()
            else:
                lblMensaje.config(text="Los campos no deben estar vacíos", fg="red")

        def actualizar():
            if validar():
                id = id_objeto.get()  # Tomar el ID del campo de texto
                if id:
                    sql = "UPDATE objeto SET nombreObjeto=%s WHERE id=%s"
                    try:
                        db.cursor.execute(sql, (nombre_objeto.get(), id))
                        db.connection.commit()
                        lblMensaje.config(text="Se ha actualizado el registro con éxito", fg="green")
                        llenar_tabla()
                        limpiar()
                    except pymysql.MySQLError as e:
                        lblMensaje.config(text=f"Error al actualizar el registro: {e}", fg="red")

        def limpiar():
            id_objeto.set("")
            nombre_objeto.set("")
            filtro_nombre.set("")

        def validar():
            return len(nombre_objeto.get()) > 0

        def vaciar_tabla():
            filas = tvObjetos.get_children()
            for fila in filas:
                tvObjetos.delete(fila)

        def llenar_tabla():
            vaciar_tabla()
            global total_records
            offset = (page_number - 1) * records_per_page
            sql_count = "SELECT COUNT(*) FROM objeto WHERE nombreObjeto LIKE %s"
            try:
                db.cursor.execute(sql_count, ('%' + filtro_nombre.get() + '%',))
                total_records = db.cursor.fetchone()[0]

                sql = "SELECT id, nombreObjeto FROM objeto WHERE nombreObjeto LIKE %s LIMIT %s OFFSET %s"
                db.cursor.execute(sql, ('%' + filtro_nombre.get() + '%', records_per_page, offset))
                filas = db.cursor.fetchall()
                for fila in filas:
                    id = fila[0]
                    tvObjetos.insert("", 'end', id, text=id, values=(fila[0], fila[1]))
            except pymysql.MySQLError as e:
                lblMensaje.config(text=f"Error al cargar datos: {e}", fg="red")
            # Actualizar el estado de los botones de navegación
            btnPrev.config(state='normal' if page_number > 1 else 'disabled')
            btnNext.config(state='normal' if page_number * records_per_page < total_records else 'disabled')

        def paginacion():
            llenar_tabla()
            # Actualizar el estado de los botones de navegación
            btnPrev.config(state='normal' if page_number > 1 else 'disabled')
            btnNext.config(state='normal' if page_number * records_per_page < total_records else 'disabled')

        def prev_page():
            global page_number
            if page_number > 1:
                page_number -= 1
                paginacion()

        def next_page():
            global page_number
            if page_number * records_per_page < total_records:
                page_number += 1
                paginacion()

        # Crear el marco de la pantalla
        marco = tk.LabelFrame(pantalla_objeto, text="Formulario Objeto")
        marco.place(x=50, y=50, width=1500, height=1400)

        # Botones para cerrar y minimizar
        btnCerrar = tk.Button(marco, text="X", command=pantalla_objeto.destroy, bg="red", fg="white")
        btnCerrar.place(x=470, y=10, width=20, height=20)

        btnMinimizar = tk.Button(marco, text="-", command=lambda: pantalla_objeto.iconify(), bg="yellow", fg="black")
        btnMinimizar.place(x=450, y=10, width=20, height=20)

        tk.Label(marco, text="ID").grid(column=0, row=0, padx=5, pady=5)
        txtIdObjeto = tk.Entry(marco, textvariable=id_objeto, state='disabled')
        txtIdObjeto.grid(column=1, row=0)

        tk.Label(marco, text="Nombre").grid(column=0, row=1, padx=5, pady=5)
        txtNombreObjeto = tk.Entry(marco, textvariable=nombre_objeto)
        txtNombreObjeto.grid(column=1, row=1)

        tk.Label(marco, text="Filtro Nombre").grid(column=0, row=2, padx=5, pady=5)
        txtFiltroNombre = tk.Entry(marco, textvariable=filtro_nombre)
        txtFiltroNombre.grid(column=1, row=2)
        filtro_nombre.trace_add('write', aplicar_filtro)

        lblMensaje = tk.Label(marco, text="Aquí van los mensajes", fg="green")
        lblMensaje.grid(column=0, row=3, columnspan=4)

        tvObjetos = ttk.Treeview(marco, selectmode='browse')  # Cambiar a modo 'browse'
        tvObjetos["columns"] = ("ID", "Nombre")
        tvObjetos.column("#0", width=0, stretch='no')
        tvObjetos.column("ID", width=150, anchor='center')
        tvObjetos.column("Nombre", width=150, anchor='center')
        tvObjetos.heading("#0", text="")
        tvObjetos.heading("ID", text="ID", anchor='center')
        tvObjetos.heading("Nombre", text="Nombre", anchor='center')
        tvObjetos.grid(column=0, row=4, columnspan=4, padx=5)
        tvObjetos.bind("<<TreeviewSelect>>", seleccionar)

        # Botones de acción
        btnEliminar = tk.Button(marco, text="Eliminar", command=eliminar)
        btnEliminar.grid(column=1, row=6, padx=5, pady=10)

        btnNuevo = tk.Button(marco, text="Guardar", command=nuevo)
        btnNuevo.grid(column=2, row=6, padx=5, pady=10)

        btnActualizar = tk.Button(marco, text="Actualizar", command=actualizar)
        btnActualizar.grid(column=3, row=6, padx=5, pady=10)

        btnLimpiar = tk.Button(marco, text="Limpiar", command=limpiar)
        btnLimpiar.grid(column=0, row=6, padx=5, pady=10)

        # Botones de navegación
        btnPrev = tk.Button(marco, text="<< Anterior", command=prev_page)
        btnPrev.grid(column=0, row=7, pady=10, sticky='w')

        btnNext = tk.Button(marco, text="Siguiente >>", command=next_page)
        btnNext.grid(column=3, row=7, pady=10, sticky='e')

        llenar_tabla()

    def abrir_obligaciones_window():
        # Crear una nueva pantalla para "Obligaciones Específicas"
        pantalla_obligaciones = tk.Toplevel()
        pantalla_obligaciones.title("CRUD Obligaciones Específicas")

        # Inicializar variables
        db = DataBase()
        global page_number, records_per_page, total_records
        page_number = 1
        records_per_page = 10
        total_records = 0

        modificar = False
        id_obligacion = tk.StringVar()
        nombre_obligacion = tk.StringVar()
        id_objeto = tk.StringVar()
        año = tk.StringVar()
        filtro_nombre = tk.StringVar()
        filtro_año = tk.StringVar()
        filtro_id_objeto = tk.StringVar()

        def aplicar_filtro(*args):
            llenar_tabla()

        def seleccionar(event):
            seleccion = tvObligaciones.selection()
            if seleccion:
                id = seleccion[0]
                valores = tvObligaciones.item(id, "values")
                id_obligacion.set(valores[0])
                nombre_obligacion.set(valores[1])
                id_objeto.set(valores[2])
                año.set(valores[3])
                btnEliminar.config(state='normal')

        def eliminar():
            seleccion = tvObligaciones.selection()
            if seleccion:
                id = seleccion[0]
                if int(id) > 0:
                    sql = "DELETE FROM obligacionesespecificas WHERE id=%s"
                    try:
                        db.cursor.execute(sql, (id,))
                        db.connection.commit()
                        tvObligaciones.delete(id)
                        lblMensaje.config(text="Se ha eliminado el registro correctamente", fg="green")
                        limpiar()
                        llenar_tabla()
                    except pymysql.MySQLError as e:
                        lblMensaje.config(text=f"Error al eliminar el registro: {e}", fg="red")
                else:
                    lblMensaje.config(text="Seleccione un registro para eliminar", fg="red")

        def nuevo():
            if not modificar:
                if validar():
                    sql_get_id = "SELECT IFNULL(MAX(id), 0) + 1 FROM obligacionesespecificas"
                    db.cursor.execute(sql_get_id)
                    nuevo_id = db.cursor.fetchone()[0]

                    sql = "INSERT INTO obligacionesespecificas (id, nombre, idObjeto, año) VALUES (%s, %s, %s, %s)"
                    try:
                        db.cursor.execute(sql, (nuevo_id, nombre_obligacion.get(), id_objeto.get(), año.get()))
                        db.connection.commit()
                        lblMensaje.config(text="Se ha guardado el registro con éxito", fg="green")
                        llenar_tabla()
                        limpiar()
                    except pymysql.MySQLError as e:
                        lblMensaje.config(text=f"Error al guardar el registro: {e}", fg="red")
                else:
                    lblMensaje.config(text="Los campos no deben estar vacíos", fg="red")
            else:
                modificarFalse()

        def actualizar():
            if modificar:
                if validar():
                    id = id_obligacion.get()
                    if id:
                        sql = "UPDATE obligacionesespecificas SET nombre=%s, idObjeto=%s, año=%s WHERE id=%s"
                        try:
                            db.cursor.execute(sql, (nombre_obligacion.get(), id_objeto.get(), año.get(), id))
                            db.connection.commit()
                            lblMensaje.config(text="Se ha guardado el registro con éxito", fg="green")
                            llenar_tabla()
                            limpiar()
                        except pymysql.MySQLError as e:
                            lblMensaje.config(text=f"Error al actualizar el registro: {e}", fg="red")
                else:
                    lblMensaje.config(text="Los campos no deben estar vacíos", fg="red")
            else:
                modificarTrue()

        def vaciar_tabla():
            filas = tvObligaciones.get_children()
            for fila in filas:
                tvObligaciones.delete(fila)

        def llenar_tabla():
            vaciar_tabla()
            global total_records
            offset = (page_number - 1) * records_per_page

            sql_count = """
            SELECT COUNT(*) 
            FROM obligacionesespecificas o 
            LEFT JOIN objeto ob ON o.idObjeto = ob.id 
            WHERE o.nombre LIKE %s 
            AND o.año LIKE %s 
            """
            try:
                db.cursor.execute(sql_count, ('%' + filtro_nombre.get() + '%', '%' + filtro_año.get() + '%'))
                total_records = db.cursor.fetchone()[0]

                sql = """
                SELECT o.id, o.nombre, ob.nombreObjeto, o.año 
                FROM obligacionesespecificas o 
                LEFT JOIN objeto ob ON o.idObjeto = ob.id 
                WHERE o.nombre LIKE %s 
                AND o.año LIKE %s 
                LIMIT %s OFFSET %s
                """
                db.cursor.execute(sql, ('%' + filtro_nombre.get() + '%', '%' + filtro_año.get() + '%', records_per_page, offset))
                filas = db.cursor.fetchall()
                for fila in filas:
                    id = fila[0]
                    tvObligaciones.insert("", 'end', id, text=id, values=(fila[0], fila[1], fila[2] if fila[2] else 'None', fila[3]))
            except pymysql.MySQLError as e:
                lblMensaje.config(text=f"Error al cargar datos: {e}", fg="red")
            
            btnPrev.config(state='normal' if page_number > 1 else 'disabled')
            btnNext.config(state='normal' if page_number * records_per_page < total_records else 'disabled')

        def limpiar():
            id_obligacion.set("")
            nombre_obligacion.set("")
            id_objeto.set("")
            año.set("")
            btnEliminar.config(state='disabled')

        def validar():
            return (len(nombre_obligacion.get()) > 0 and 
                    len(id_objeto.get()) > 0 and 
                    len(año.get()) > 0)

        def modificarFalse():
            nonlocal modificar
            modificar = False
            tvObligaciones.config(selectmode='none')
            btnNuevo.config(text="Guardar")
            btnActualizar.config(text="Actualizar")
            btnEliminar.config(state='disabled')

        def modificarTrue():
            nonlocal modificar
            modificar = True
            tvObligaciones.config(selectmode='browse')
            btnNuevo.config(text="Nuevo")
            btnActualizar.config(text="Actualizar")
            btnEliminar.config(state='normal')

        def paginacion():
            llenar_tabla()
            btnPrev.config(state='normal' if page_number > 1 else 'disabled')
            btnNext.config(state='normal' if page_number * records_per_page < total_records else 'disabled')

        def prev_page():
            global page_number
            if page_number > 1:
                page_number -= 1
                paginacion()

        def next_page():
            global page_number
            if page_number * records_per_page < total_records:
                page_number += 1
                paginacion()

        def cargar_objetos():
            sql = "SELECT id, nombreObjeto FROM objeto"
            try:
                db.cursor.execute(sql)
                objetos = db.cursor.fetchall()
                comboIdObjeto['values'] = [obj[1] for obj in objetos]
            except pymysql.MySQLError as e:
                lblMensaje.config(text=f"Error al cargar objetos: {e}", fg="red")

        # Crear el marco de la pantalla
        marco = tk.LabelFrame(pantalla_obligaciones, text="Formulario Obligaciones Específicas")
        marco.place(x=50, y=50, width=1500, height=700)

        # Botones para cerrar y minimizar
        btnCerrar = tk.Button(marco, text="X", command=pantalla_obligaciones.destroy, bg="red", fg="white")
        btnCerrar.place(x=470, y=10, width=20, height=20)

        btnMinimizar = tk.Button(marco, text="-", command=lambda: pantalla_obligaciones.iconify(), bg="yellow", fg="black")
        btnMinimizar.place(x=450, y=10, width=20, height=20)

        # Crear marco para la sección de registro
        marco_registro = tk.Frame(marco)
        marco_registro.grid(row=0, column=0, padx=10, pady=10)

        # Etiquetas y campos de entrada
        tk.Label(marco_registro, text="ID Obligación").grid(row=0, column=0, padx=5, pady=5)
        tk.Entry(marco_registro, textvariable=id_obligacion, state='disabled').grid(row=0, column=1, padx=5, pady=5)

        tk.Label(marco_registro, text="Nombre").grid(row=1, column=0, padx=5, pady=5)
        tk.Entry(marco_registro, textvariable=nombre_obligacion).grid(row=1, column=1, padx=5, pady=5)

        tk.Label(marco_registro, text="ID Objeto").grid(row=2, column=0, padx=5, pady=5)

        comboIdObjeto = ttk.Combobox(marco_registro, textvariable=id_objeto)
        comboIdObjeto.grid(row=2, column=1, padx=5, pady=5)
        comboIdObjeto.bind("<<ComboboxSelected>>", lambda e: id_objeto.set(comboIdObjeto.get()))
        cargar_objetos()

        tk.Label(marco_registro, text="Año").grid(row=3, column=0, padx=5, pady=5)
        tk.Entry(marco_registro, textvariable=año).grid(row=3, column=1, padx=5, pady=5)

        # Botones para guardar, actualizar y eliminar
        btnNuevo = tk.Button(marco_registro, text="Guardar", command=nuevo)
        btnNuevo.grid(row=4, column=0, padx=5, pady=5)

        btnActualizar = tk.Button(marco_registro, text="Actualizar", command=actualizar)
        btnActualizar.grid(row=4, column=1, padx=5, pady=5)

        btnEliminar = tk.Button(marco_registro, text="Eliminar", command=eliminar, state='disabled')
        btnEliminar.grid(row=4, column=2, padx=5, pady=5)

        # Mensaje
        lblMensaje = tk.Label(marco_registro, text="", fg="red")
        lblMensaje.grid(row=5, column=0, columnspan=3)

        # Crear marco para la sección de filtros
        marco_filtros = tk.Frame(marco)
        marco_filtros.grid(row=0, column=1, padx=10, pady=10)

        # Etiquetas y campos de filtro
        tk.Label(marco_filtros, text="Nombre").grid(row=0, column=0, padx=5, pady=5)
        tk.Entry(marco_filtros, textvariable=filtro_nombre).grid(row=0, column=1, padx=5, pady=5)

        tk.Label(marco_filtros, text="Año").grid(row=1, column=0, padx=5, pady=5)
        tk.Entry(marco_filtros, textvariable=filtro_año).grid(row=1, column=1, padx=5, pady=5)

        tk.Label(marco_filtros, text="ID Objeto").grid(row=2, column=0, padx=5, pady=5)
        tk.Entry(marco_filtros, textvariable=filtro_id_objeto).grid(row=2, column=1, padx=5, pady=5)

        # Botón para aplicar filtro
        btnFiltrar = tk.Button(marco_filtros, text="Filtrar", command=aplicar_filtro)
        btnFiltrar.grid(row=3, column=0, columnspan=2)

        # Crear la tabla
        columnas = ("ID", "Nombre", "ID Objeto", "Año")
        tvObligaciones = ttk.Treeview(marco, columns=columnas, show="headings")
        tvObligaciones.grid(row=1, column=0, columnspan=2, padx=10, pady=10)

        for col in columnas:
            tvObligaciones.heading(col, text=col)

        tvObligaciones.bind("<ButtonRelease-1>", seleccionar)

        # Crear botones de paginación
        btnPrev = tk.Button(marco, text="<<", command=prev_page)
        btnPrev.grid(row=2, column=0, sticky='w')

        btnNext = tk.Button(marco, text=">>", command=next_page)
        btnNext.grid(row=2, column=1, sticky='e')

        # Llenar tabla inicialmente
        llenar_tabla()

    def validar_numeros(char):
        return char.isdigit()  # Permitir solo caracteres numéricos
    # Función para obtener objetos desde la base de datos
    def obtener_objetos():
        try:
            # Conexión a la base de datos
            conexion = pymysql.connect(host='localhost', user='root', password='', db='contrataciones')
            cursor = conexion.cursor()

            # Consulta SQL para obtener los nombres de los objetos
            cursor.execute("SELECT nombre FROM objetos")
            resultados = cursor.fetchall()

            # Extraer los nombres de los objetos y almacenarlos en una lista
            nombres_objetos = [fila[0] for fila in resultados]

            conexion.close()  # Cerrar la conexión

            return nombres_objetos

        except Exception as e:
            print(f"Error al obtener objetos: {e}")
            return []

    # Cargar los nombres de objetos en la variable
    nombres_objetos = obtener_objetos()
    # Función para obtener las obligaciones específicas desde la base de datos
    def obtener_obligaciones_especificas():
        try:
            # Conexión a la base de datos
            conexion = pymysql.connect(host='localhost', user='root', password='', db='contrataciones')
            cursor = conexion.cursor()

            # Consulta SQL para obtener las obligaciones específicas
            cursor.execute("SELECT nombre FROM obligaciones_especificas")
            resultados = cursor.fetchall()

            # Extraer los nombres de las obligaciones y almacenarlos en una lista
            nombres_obligaciones_especificas = [fila[0] for fila in resultados]

            conexion.close()  # Cerrar la conexión

            return nombres_obligaciones_especificas

        except Exception as e:
            print(f"Error al obtener obligaciones específicas: {e}")
            return []

    # Cargar los nombres de obligaciones en la variable
    nombres_obligaciones_especificas = obtener_obligaciones_especificas()
    # Crear el validador numérico
    validacion_numerica = contenedor_formulario.register(validar_numeros)

    # Filas para organizar los inputs y selects
    for i in range(5):  # Cambiar el rango a 5 para incluir la fila adicional para "Objeto"
        contenedor_fila = tk.Frame(contenedor_formulario)
        contenedor_fila.pack(fill='x', pady=5)
        if i == 0:
            # Fila 1
            tk.Label(contenedor_fila, text="ID Independiente:").grid(row=0, column=0, sticky="e")
            entry_id_independiente = tk.Entry(contenedor_fila, textvariable=id_independiente, state='readonly')
            entry_id_independiente.grid(row=0, column=1, sticky="ew")

            # Alineando los siguientes elementos en la misma fila
            tk.Label(contenedor_fila, text="Cliente:").grid(row=0, column=2, sticky="e")
            comboClientes = ttk.Combobox(contenedor_fila)
            comboClientes.grid(row=0, column=3, sticky="ew")

            tk.Label(contenedor_fila, text="Tipo de Contrato:").grid(row=0, column=4, sticky="e")
            comboTipoContrato = ttk.Combobox(contenedor_fila)
            comboTipoContrato.grid(row=0, column=5, sticky="ew")

            tk.Label(contenedor_fila, text="Cargo:").grid(row=0, column=6, sticky="e")
            comboCargo = ttk.Combobox(contenedor_fila)
            comboCargo.grid(row=0, column=7, sticky="ew")

            tk.Label(contenedor_fila, text="Dependencia:").grid(row=0, column=8, sticky="e")
            comboDependencia = ttk.Combobox(contenedor_fila)
            comboDependencia.grid(row=0, column=9, sticky="ew")


        elif i == 1:
            # Fila 1 - Descripción y otros elementos
            tk.Label(contenedor_fila, text="Descripción:").grid(row=1, column=0, sticky="e")
            tk.Entry(contenedor_fila, textvariable=descripcion_contrato).grid(row=1, column=1, sticky="ew")

            tk.Label(contenedor_fila, text="Vigencia:").grid(row=1, column=2, sticky="e")
            date_vigencia = DateEntry(contenedor_fila, width=20, background='darkblue',
                                    foreground='white', borderwidth=2, date_pattern='y-mm-dd')
            date_vigencia.grid(row=1, column=3, sticky="ew")

            tk.Label(contenedor_fila, text="Terminación:").grid(row=1, column=4, sticky="e")
            date_terminacion = DateEntry(contenedor_fila, width=20, background='darkblue',
                                        foreground='white', borderwidth=2, date_pattern='y-mm-dd')
            date_terminacion.grid(row=1, column=5, sticky="ew")

            # Fila 2 - Autorización y Valor (se alinean en la misma fila)
            tk.Label(contenedor_fila, text="Autorización:").grid(row=1, column=6, sticky="e")
            tk.Entry(contenedor_fila, textvariable=autorizacion_contratos, validate="key", 
                    validatecommand=(validacion_numerica, '%S')).grid(row=1, column=7, sticky="ew")

            tk.Label(contenedor_fila, text="Valor:").grid(row=1, column=8, sticky="e")
            tk.Entry(contenedor_fila, textvariable=valor_contrato, validate="key", 
                    validatecommand=(validacion_numerica, '%S')).grid(row=1, column=9, sticky="ew")


        elif i == 2:  # Nueva fila para "Objeto"
            tk.Label(contenedor_fila, text="Objeto:").grid(row=0, column=0, sticky="e")
            comboObjeto = ttk.Combobox(contenedor_fila, textvariable=id_version, values=nombres_objetos, width=40)  # Adjust 'width' to desired value
            comboObjeto.grid(row=0, column=1, sticky="ew", ipady=5)
            comboObjeto.bind("<KeyRelease>", obtener_objetos)  # Filtrar mientras escribes

            tk.Label(contenedor_fila, text="Obligaciones Específicas:").grid(row=0, column=2, sticky="e")
            comboObligaciones = ttk.Combobox(contenedor_fila, textvariable=obligaciones_especificas, values=nombres_obligaciones_especificas, width=40,  justify="left")
            comboObligaciones.grid(row=0, column=3, columnspan=3, sticky="ew", ipady=5)
            comboObligaciones.bind("<KeyRelease>", obtener_obligaciones_especificas)  # Filtrar mientras escribes

            tk.Label(contenedor_fila, text="Consecutivo:").grid(row=0, column=6, sticky="e")
            entry_consecutivo = tk.Entry(contenedor_fila, textvariable=consecutivo, validate="key", 
                                        validatecommand=(validacion_numerica, '%S'))
            entry_consecutivo.grid(row=0, column=7, sticky="ew")

            tk.Label(contenedor_fila, text="Consecutivo:").grid(row=0, column=6, sticky="e")
            entry_consecutivo = tk.Entry(contenedor_fila, textvariable=consecutivo, validate="key", 
                                        validatecommand=(validacion_numerica, '%S'))
            entry_consecutivo.grid(row=0, column=7, sticky="ew")

            # Botón para crear nuevo objeto
            tk.Button(contenedor_fila, text="Crear Objeto", command=abrir_objeto_window).grid(row=0, column=8, padx=5)

            # Botón para crear nuevas obligaciones específicas
            tk.Button(contenedor_fila, text="Crear Obligaciones", command=abrir_obligaciones_window).grid(row=0, column=9, padx=5)
            
    # Botones
    btnGuardar = tk.Button(contenedor_formulario, text="Guardar", command=guardar_contrato)
    btnGuardar.pack(side='left', padx=5, pady=5)

    btnEliminar = tk.Button(contenedor_formulario, text="Eliminar", command=eliminar_contrato)
    btnEliminar.pack(side='left', padx=5, pady=5)

    btnLimpiar = tk.Button(contenedor_formulario, text="Limpiar", command=limpiar_campos)
    btnLimpiar.pack(side='left', padx=5, pady=5)

    # Campo de entrada para el filtro de `idVersion`
    tk.Label(contenedor_formulario, text="Filtro ID Versión:").pack(side='left', padx=5, pady=5)
    entry_filtro_id_version = tk.Entry(contenedor_formulario, textvariable=filtro_id_version)
    entry_filtro_id_version.pack(side='left', padx=5, pady=5)

    # Asociar la actualización automática al cambio de valor del filtro
    filtro_id_version.trace_add("write", lambda *args: llenar_tabla())

    # Scrollbar horizontal para la tabla
    scroll_horizontal = tk.Scrollbar(marco, orient='horizontal')
    scroll_horizontal.pack(side='bottom', fill='x')

    # Tabla (Treeview) con scroll horizontal
    tabla = ttk.Treeview(marco, columns=('ID', 'idVersion', 'Numero de documento', 'Cliente', 'Tipo Contrato', 'Cargo', 'Dependencia', 'Objeto', 'Descripción', 'Vigencia', 'Terminación', 'Autorización', 'Valor', 'obligaciones_especificas', 'consecutivo'), show='headings', xscrollcommand=scroll_horizontal.set)
    tabla.pack(fill='both', expand=True, padx=10, pady=10)

    # Configurar el scrollbar para que funcione con la tabla
    scroll_horizontal.config(command=tabla.xview)

    # Configurar las cabeceras de la tabla
    for col in tabla['columns']:
        tabla.heading(col, text=col)

    tabla.bind('<<TreeviewSelect>>', seleccionar_contrato)

    # Botones de Paginación
    contenedor_paginacion = tk.Frame(marco)
    contenedor_paginacion.pack(fill='x', padx=10, pady=5)

    btnAnterior = tk.Button(contenedor_paginacion, text="<< Anterior", command=ir_a_pagina_anterior)
    btnAnterior.pack(side='left', fill='x', expand=True)

    btnSiguiente = tk.Button(contenedor_paginacion, text="Siguiente >>", command=ir_a_pagina_siguiente)
    btnSiguiente.pack(side='right', fill='x', expand=True)

    # Campos de filtro
    contenedor_filtros = tk.Frame(marco)
    contenedor_filtros.pack(fill='x', padx=10, pady=5)

    tk.Label(contenedor_filtros, text="Filtro Cliente:").pack(side='left')
    comboFiltroCliente = ttk.Combobox(contenedor_filtros, textvariable=filtro_cliente)
    comboFiltroCliente.pack(side='left', fill='x', expand=True)
    filtro_cliente.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Tipo de Contrato:").pack(side='left')
    comboFiltroTipoContrato = ttk.Combobox(contenedor_filtros, textvariable=filtro_tipo_contrato)
    comboFiltroTipoContrato.pack(side='left', fill='x', expand=True)
    filtro_tipo_contrato.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Cargo:").pack(side='left')
    comboFiltroCargo = ttk.Combobox(contenedor_filtros, textvariable=filtro_cargo)
    comboFiltroCargo.pack(side='left', fill='x', expand=True)
    filtro_cargo.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Dependencia:").pack(side='left')
    comboFiltroDependencia = ttk.Combobox(contenedor_filtros, textvariable=filtro_dependencia)
    comboFiltroDependencia.pack(side='left', fill='x', expand=True)
    filtro_dependencia.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Autorización:").pack(side='left')
    tk.Entry(contenedor_filtros, textvariable=filtro_autorizacion).pack(side='left', fill='x', expand=True)
    filtro_autorizacion.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Valor:").pack(side='left')
    tk.Entry(contenedor_filtros, textvariable=filtro_valor).pack(side='left', fill='x', expand=True)
    filtro_valor.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Objeto:").pack(side='left')
    comboFiltroObjeto = ttk.Combobox(contenedor_filtros, textvariable=filtro_id_comboObjeto)
    comboFiltroObjeto.pack(side='left', fill='x', expand=True)
    filtro_id_comboObjeto.trace_add("write", lambda *args: llenar_tabla())

    tk.Label(contenedor_filtros, text="Filtro Obligaciones Específicas:").pack(side='left')
    comboFiltroObligaciones_Específicas  = ttk.Combobox(contenedor_filtros, textvariable=filtro_ObEspCon)
    comboFiltroObligaciones_Específicas.pack(side='left', fill='x', expand=True)
    filtro_id_Obligaciones.trace_add("write", lambda *args: llenar_tabla())

    # Botón de Mostrar Todos los Registros
    btnMostrarTodos = tk.Button(marco, text="Mostrar Todos los Registros", command=mostrar_todos_los_registros)
    btnMostrarTodos.pack(fill='x', padx=10, pady=5)

    # Función para calcular el término de ejecución
    def calcular_diferencia_fechas(fecha_inicio, fecha_fin):
        # Convertir las fechas de cadena a objetos datetime
        inicio = datetime.strptime(fecha_inicio, '%Y-%m-%d')
        fin = datetime.strptime(fecha_fin, '%Y-%m-%d')

        # Calcular la diferencia en meses y días
        diferencia = relativedelta(fin, inicio)
        meses = diferencia.years * 12 + diferencia.months
        dias = diferencia.days

        # Si no ha pasado al menos un mes completo, dejar meses en 0
        if meses == 0 and dias > 0:
            meses = 0

        # Formatear en texto con números entre paréntesis
        resultado = f"{meses} Meses y {dias} días"
        return resultado

  # Función para convertir número a texto
    def convertir_valor_a_texto(valor):
        valor_texto = num2words(valor, lang='es').upper()  # Convertimos el número a palabras en español y en mayúsculas
        return f"{valor_texto} PESOS M/CTE."
    # Función para exportar a Word
    # Función para obtener la fecha actual en el formato requerido
    def obtener_fecha_actual_formateada():
        # Obtener la fecha actual
        fecha_actual = datetime.now()

        # Día en número
        dia = fecha_actual.day

        # Mes en texto
        mes = fecha_actual.strftime('%B').capitalize()

        # Año en número
        año = fecha_actual.year

        # Formato final: a los (día en número) días del mes de (mes en texto) del año (año en número)
        return f"a los {dia} días del mes de {mes} del año {año}"

    # Función para obtener el id, nombre completo y la ciudad de expedición del cliente
    def obtener_datos_cliente(cliente_id):
        try:
            # Consultar la información del cliente: id, nombres, apellidos y ciudad de expedición
            sql = """
            SELECT id, primerNombre, segundoNombre, primerApellido, segundoApellido, ciudadExpedicion
            FROM clientes
            WHERE id = %s
            """
            db.cursor.execute(sql, (cliente_id,))
            datos_cliente = db.cursor.fetchone()

            if datos_cliente:
                # Concatenamos los nombres y apellidos para formar el nombre completo
                nombre_completo = f"{datos_cliente[1]} {datos_cliente[2]} {datos_cliente[3]} {datos_cliente[4]}"
                ciudad_expedicion = datos_cliente[5]
                return datos_cliente[0], nombre_completo, ciudad_expedicion
            else:
                print(f"No se encontraron datos para el cliente con id {cliente_id}.")
                return None, None, None

        except Exception as e:
            print(f"Error al obtener datos del cliente: {e}")
            return None, None, None
    # Función para obtener el id del cliente basado en el nombre completo
    def obtener_id_cliente_por_nombre(nombre_completo):
        try:
            # Separar el nombre completo en las partes correspondientes
            partes_nombre = nombre_completo.split()
            primer_nombre = partes_nombre[0]
            segundo_nombre = partes_nombre[1] if len(partes_nombre) > 1 else ""
            primer_apellido = partes_nombre[2] if len(partes_nombre) > 2 else ""
            segundo_apellido = partes_nombre[3] if len(partes_nombre) > 3 else ""

            # Consulta SQL para obtener el id del cliente basado en su nombre completo
            sql = """
            SELECT id, ciudadExpedicion FROM clientes
            WHERE primerNombre = %s AND segundoNombre = %s AND primerApellido = %s AND segundoApellido = %s
            """
            print(f"Ejecutando consulta SQL para obtener ID: {sql} con nombre: {nombre_completo}")
            db.cursor.execute(sql, (primer_nombre, segundo_nombre, primer_apellido, segundo_apellido))
            resultado = db.cursor.fetchone()

            if resultado:
                cliente_id = resultado[0]
                ciudad_expedicion = resultado[1]  # Obtenemos el id de la ciudadExpedicion
                print(f"ID del cliente obtenido: {cliente_id}, CiudadExpedicion ID: {ciudad_expedicion}")

                # Consulta SQL para obtener el nombre de la ciudad basado en el id de ciudadExpedicion
                sql_ciudad = """
                SELECT nombreCiudad FROM ciudad
                WHERE id = %s
                """
                db.cursor.execute(sql_ciudad, (ciudad_expedicion,))
                resultado_ciudad = db.cursor.fetchone()

                if resultado_ciudad:
                    nombreCiudad = resultado_ciudad[0]  # Obtenemos el nombre de la ciudad
                    print(f"Nombre de la ciudad obtenido: {nombreCiudad}")
                    return cliente_id, nombreCiudad
                else:
                    print(f"No se encontró la ciudad con id {ciudad_expedicion}")
                    return cliente_id, None
            else:
                print(f"No se encontró cliente con el nombre {nombre_completo}")
                return None, None

        except Exception as e:
            print(f"Error al obtener ID del cliente y la ciudad: {e}")
            return None, None
    # Función para obtener datos completos del cliente y contrato
    def obtener_datos_completos_cliente_y_contrato(cliente_id):
        try:
            sql = """
            SELECT 
                clientes.id, 
                CONCAT_WS(' ', clientes.primerNombre, clientes.segundoNombre, clientes.primerApellido, clientes.segundoApellido) AS nombreCompleto, 
                clientes.ciudadExpedicion, 
                contrato.ObEspCon, 
                contrato.autorizacionContratos,
                obligacionesespecificas.nombre AS nombreObligacion  -- Obtener el nombre de las obligaciones
            FROM 
                clientes
            JOIN 
                contrato ON clientes.id = contrato.idClientes
            LEFT JOIN 
                obligacionesespecificas ON contrato.ObEspCon = obligacionesespecificas.id  -- Unión para obtener el nombre
            WHERE 
                clientes.id = %s
            LIMIT 1
            """
            db.cursor.execute(sql, (cliente_id,))
            resultado = db.cursor.fetchone()
            if resultado:
                cliente_id, nombre_completo, ciudad_expedicion, obligaciones_especificas_id, autorizacion_contratos, nombre_obligacion = resultado
                return cliente_id, nombre_completo, ciudad_expedicion, nombre_obligacion, autorizacion_contratos  # Cambia para devolver el nombre
            else:
                return None, None, None, None, None
        except Exception as e:
            print(f"Error al obtener los datos del cliente y obligaciones: {e}")
            return None, None, None, None, None
    def nueva_version():
        try:
            # Obtener el valor actual del ID
            id_actual = id_independiente.get()

            if not id_actual:
                print("Error: No hay un ID seleccionado.")
                return

            # Consultar cuántos registros derivados existen para el ID actual
            sql = "SELECT COUNT(*) FROM contrato WHERE idVersion LIKE %s"
            db.cursor.execute(sql, (f"{id_actual}-%",))  # Usar el formato de idVersion derivado
            count_derivados = db.cursor.fetchone()[0]

            # Crear el valor para `idVersion` como `id-<número siguiente>`
            nuevo_id_version = f"{id_actual}-{count_derivados + 1}"

            # Obtener los datos del contrato actual
            sql_select = """
            SELECT idClientes, idTipoContrato, idCargo, idDependecia, idObjeto, 
            descripcionContrato, Vigencia, terminacion, autorizacionContratos, valorContrato, 
            ObEspCon, consecutivo FROM contrato WHERE id = %s
            """
            db.cursor.execute(sql_select, (id_actual,))
            contrato = db.cursor.fetchone()

            if contrato:
                # Desempaquetar los valores del contrato
                id_clientes, id_tipo_contrato, id_cargo, id_dependencia, id_idObjeto, descripcion, vigencia, terminacion, autorizacion, valor, obligaciones_especificas, consecutivo = contrato

                # Verificar la existencia de claves foráneas en las tablas relacionadas
                def verificar_existencia(tabla, campo, valor):
                    sql_verificar = f"SELECT id FROM {tabla} WHERE {campo} = %s"
                    db.cursor.execute(sql_verificar, (valor,))
                    return db.cursor.fetchone() is not None

                # Validar que todos los valores relacionados (claves foráneas) existan
                if (verificar_existencia("clientes", "id", id_clientes) and
                    verificar_existencia("tipodecontrato", "id", id_tipo_contrato) and
                    verificar_existencia("cargo", "id", id_cargo) and
                    verificar_existencia("dependencia", "id", id_dependencia) and
                    verificar_existencia("obligacionesespecificas", "id", obligaciones_especificas) and
                    verificar_existencia("objeto", "id", id_idObjeto)):

                    # Si todas las claves foráneas son válidas, insertar el nuevo registro
                    sql_insert = """
                    INSERT INTO contrato 
                    (idClientes, idTipoContrato, idCargo, idDependecia, idObjeto, descripcionContrato, 
                    Vigencia, terminacion, autorizacionContratos, valorContrato, ObEspCon, consecutivo, idVersion) 
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """
                    valores = (
                        id_clientes, id_tipo_contrato, id_cargo, id_dependencia, id_idObjeto, descripcion,
                        vigencia, terminacion, autorizacion, valor, obligaciones_especificas,
                        consecutivo, nuevo_id_version
                    )
                    db.cursor.execute(sql_insert, valores)
                    db.cursor.connection.commit()

                    print(f"Nueva versión del contrato guardada con idVersion: {nuevo_id_version}")

                    # Actualizar la tabla en la interfaz
                    llenar_tabla()

                else:
                    print("Error: No se encontraron todas las claves foráneas requeridas en las tablas relacionadas.")

            else:
                print("Error: No se encontró el contrato con el ID proporcionado.")

        except Exception as e:
            print(f"Error al generar nueva versión del contrato: {e}")
            db.cursor.connection.rollback()

    # Función para exportar a Word
    def exportar_a_word():
        try:
            selected_item = tabla.selection()
            if selected_item:
                item = tabla.item(selected_item)
                contrato_seleccionado = item['values']
                objeto = contrato_seleccionado[7]
                obligaciones_especificas = comboObligaciones.get()  # Obtiene las obligaciones como un solo texto

                nombre_completo = contrato_seleccionado[3]
                consecutivo = contrato_seleccionado[14]
                print(f"Valor de consecutivo obtenido: {consecutivo}")

                cliente_id, nombreCiudad = obtener_id_cliente_por_nombre(nombre_completo)
                if cliente_id is None or nombreCiudad is None:
                    print(f"Error: No se pudo obtener el ID del cliente o la ciudad de expedición para {nombre_completo}.")
                    return

                plantilla_path = "Certificación Contractual No. 16X PLANTILLA DE CERTIFICACIONES.docx"
                doc = Document(plantilla_path)

                cliente_id, nombre_completo, ciudad_expedicion, nombre_obligacion, autorizacion_contratos = obtener_datos_completos_cliente_y_contrato(cliente_id)
                if cliente_id is None or nombre_completo is None or ciudad_expedicion is None or nombre_obligacion is None or autorizacion_contratos is None:
                    print("Error: No se pudo obtener la información del cliente.")
                    return

                vigencia_formateada = formatear_fecha_larga(datetime.strptime(str(contrato_seleccionado[9]), '%Y-%m-%d'))
                terminacion_formateada = formatear_fecha_larga(datetime.strptime(str(contrato_seleccionado[10]), '%Y-%m-%d'))
                termino_ejecutado = calcular_diferencia_fechas_texto(str(contrato_seleccionado[9]), str(contrato_seleccionado[10]))
                valor_contrato_texto = convertir_valor_a_texto(int(contrato_seleccionado[11]))
                fecha_actual_formateada = obtener_fecha_actual_formateada()

                for paragraph in doc.paragraphs:
                    if "XXXXXXXXXXXXX" in paragraph.text:
                        paragraph.text = paragraph.text.replace("XXXXXXXXXXXXX", nombre_completo)
                    if "XXXXXXX" in paragraph.text:
                        paragraph.text = paragraph.text.replace("XXXXXXX", str(cliente_id))
                    if "CO1.PCCNTR.5780050" in paragraph.text:
                        paragraph.text = paragraph.text.replace("CO1.PCCNTR.5780050", f"CO1.PCCNTR.{autorizacion_contratos}")
                    if "47.930.183" in paragraph.text:
                        paragraph.text = paragraph.text.replace("47.930.183", str(contrato_seleccionado[11]))
                    if "CUARENTA Y SIETE MILLONES NOVECIENTOS TREINTA MIL CIENTO OCHENTA Y TRES PESOS M/CTE." in paragraph.text:
                        paragraph.text = paragraph.text.replace("CUARENTA Y SIETE MILLONES NOVECIENTOS TREINTA MIL CIENTO OCHENTA Y TRES PESOS M/CTE.", valor_contrato_texto)
                    if "22 de enero de 2024" in paragraph.text:
                        paragraph.text = paragraph.text.replace("22 de enero de 2024", vigencia_formateada)
                    if "texto-objeto" in paragraph.text:
                        paragraph.text = paragraph.text.replace("texto-objeto", objeto)
                    if "13 de diciembre de 2024" in paragraph.text:
                        paragraph.text = paragraph.text.replace("13 de diciembre de 2024", terminacion_formateada)
                    if "Diez (10) Meses y Veintidós (22) días" in paragraph.text:
                        paragraph.text = paragraph.text.replace("Diez (10) Meses y Veintidós (22) días", termino_ejecutado)
                    if "a los 12 días del mes de julio del año 2024" in paragraph.text:
                        paragraph.text = paragraph.text.replace("a los 12 días del mes de julio del año 2024", fecha_actual_formateada)
                    if "valor Ciudad" in paragraph.text:
                        paragraph.text = paragraph.text.replace("valor Ciudad", nombreCiudad)

                    for paragraph in doc.paragraphs:
                        if "Obligaciones Específicas del Contrato: " in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run("Obligaciones Específicas del Contrato: ")
                            run.bold = True
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                            run = paragraph.add_run(obligaciones_especificas.strip())
                            run.font.name = 'Calibri'
                            run.font.size = Pt(11)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            paragraph.paragraph_format.space_before = Pt(10)
                            paragraph.paragraph_format.space_after = Pt(10)
                            paragraph.paragraph_format.line_spacing = Pt(16)

                for section in doc.sections:
                    header = section.header
                    for paragraph in header.paragraphs:
                        if "Certificación No. Consecutivo" in paragraph.text:
                            for run in paragraph.runs:
                                if "Consecutivo" in run.text:
                                    run.text = run.text.replace("Consecutivo", str(consecutivo))
                                    # Remover resaltado (resaltado amarillo)
                                    run.font.highlight_color = None
                            print(f"Encabezado actualizado con Consecutivo: {consecutivo}")
                doc_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
                if doc_path:
                    doc.save(doc_path)
                    print(f"Archivo Word guardado en: {doc_path}")
                    return doc_path  # Devolvemos la ruta para convertirla luego a PDF
                else:
                    print("Operación cancelada por el usuario.")
                    return None

        except Exception as e:
            print(f"Error al exportar a Word: {e}")
            return None
    # Función para calcular la diferencia en años, meses y días y formatearla en texto
    def calcular_diferencia_fechas_texto(fecha_inicio, fecha_fin):
        fecha_inicio_dt = datetime.strptime(fecha_inicio, '%Y-%m-%d')
        fecha_fin_dt = datetime.strptime(fecha_fin, '%Y-%m-%d')
        
        diferencia = relativedelta(fecha_fin_dt, fecha_inicio_dt)
        anios = diferencia.years
        meses = diferencia.months
        dias = diferencia.days

        anios_texto = convertir_numero_a_texto(anios)
        meses_texto = convertir_numero_a_texto(meses)
        dias_texto = convertir_numero_a_texto(dias)

        # Formato con años, meses y días
        return f"{anios_texto.capitalize()} ({anios}) año(s), {meses_texto.capitalize()} ({meses}) meses y {dias_texto.capitalize()} ({dias}) días"

    # Función auxiliar para convertir números a texto
    def convertir_numero_a_texto(numero):
        numeros_texto = {
            0: 'Cero', 1: 'Uno', 2: 'Dos', 3: 'Tres', 4: 'Cuatro', 5: 'Cinco', 6: 'Seis', 7: 'Siete', 8: 'Ocho', 9: 'Nueve', 10: 'Diez',
            11: 'Once', 12: 'Doce', 13: 'Trece', 14: 'Catorce', 15: 'Quince', 16: 'Dieciséis', 17: 'Diecisiete', 18: 'Dieciocho', 19: 'Diecinueve', 20: 'Veinte',
            21: 'Veintiuno', 22: 'Veintidos', 23: 'Veintitres', 24: 'Veinticuatro', 25: 'Veinticinco', 26: 'Veintiseis', 27: 'Veintisiete', 28: 'Veintiocho', 29: 'Veintinueve', 30: 'Treinta', 31: 'Treinta y uno'
        }
        return numeros_texto.get(numero, str(numero))
    
    # Función para exportar a PDF
    def exportar_a_pdf():
        try:
            # Primero, exportamos el archivo a Word y obtenemos su ruta
            doc_path = exportar_a_word()
            if doc_path:
                # Luego, convertimos ese archivo de Word a PDF
                pdf_path = doc_path.replace(".pdf")
                convert(pdf_path)
                print(f"Archivo PDF guardado en: {pdf_path}")
            else:
                print("No se pudo generar el PDF ya que el archivo Word no fue exportado correctamente.")
        
        except Exception as e:
            print(f"Error al exportar a PDF: {e}")
    # Función para calcular la diferencia en años, meses y días y formatearla en texto
    def calcular_diferencia_fechas_texto(fecha_inicio, fecha_fin):
        fecha_inicio_dt = datetime.strptime(fecha_inicio, '%Y-%m-%d')
        fecha_fin_dt = datetime.strptime(fecha_fin, '%Y-%m-%d')
        
        diferencia = relativedelta(fecha_fin_dt, fecha_inicio_dt)
        anios = diferencia.years
        meses = diferencia.months
        dias = diferencia.days

        anios_texto = convertir_numero_a_texto(anios)
        meses_texto = convertir_numero_a_texto(meses)
        dias_texto = convertir_numero_a_texto(dias)

        # Formato con años, meses y días
        return f"{anios_texto.capitalize()} ({anios}) año(s), {meses_texto.capitalize()} ({meses}) Meses y {dias_texto.capitalize()} ({dias}) días"

    # Función auxiliar para convertir números a texto
    def convertir_numero_a_texto(numero):
        numeros_texto = {
            0: 'Cero', 1: 'Uno', 2: 'Dos', 3: 'Tres', 4: 'Cuatro', 5: 'Cinco', 6: 'Seis', 7: 'Siete', 8: 'Ocho', 9: 'Nueve', 10: 'Diez',
            11: 'Once', 12: 'Doce', 13: 'Trece', 14: 'Catorce', 15: 'Quince', 16: 'Dieciséis', 17: 'Diecisiete', 18: 'Dieciocho', 19: 'Diecinueve', 20: 'Veinte',
            21: 'Veintiuno', 22: 'Veintidós', 23: 'Veintitrés', 24: 'Veinticuatro', 25: 'Veinticinco', 26: 'Veintiséis', 27: 'Veintisiete', 28: 'Veintiocho', 29: 'Veintinueve', 30: 'Treinta', 31: 'Treinta y uno'
        }
        return numeros_texto.get(numero, str(numero))

    def exportar_a_excel():
        try:
            # Abrir un diálogo para que el usuario seleccione la ubicación y nombre del archivo
            archivo_excel = filedialog.asksaveasfilename(
                defaultextension=".xlsx",
                filetypes=[("Archivos de Excel", "*.xlsx")],
                title="Guardar archivo de Excel"
            )
            
            if archivo_excel:  # Si el usuario selecciona una ubicación
                # Conectar a la base de datos
                conexion = pymysql.connect(host='localhost', user='root', password='', db='contrataciones')
                cursor = conexion.cursor()

                # Consulta SQL para obtener los datos con nombres en lugar de IDs y agregar idVersion e idClientes
                consulta = """
                SELECT
                    contrato.id,
                    contrato.idVersion,
                    contrato.idClientes,
                    contrato.descripcionContrato,
                    contrato.Vigencia,
                    contrato.terminacion,
                    contrato.autorizacionContratos,
                    contrato.valorContrato,
                    contrato.consecutivo,

                    CONCAT(clientes.primerNombre, ' ', clientes.segundoNombre, ' ', clientes.primerApellido, ' ', clientes.segundoApellido) AS nombreCliente,

                    tipodecontrato.nombreTipoContrato,

                    cargo.nombreCargo,

                    dependencia.nombreDependencia,

                    objeto.nombreObjeto,

                    obligacionesespecificas.nombre AS nombreObligacionEspecifica

                FROM contrato
                LEFT JOIN clientes ON contrato.idClientes = clientes.id
                LEFT JOIN tipodecontrato ON contrato.idTipoContrato = tipodecontrato.id
                LEFT JOIN cargo ON contrato.idCargo = cargo.id
                LEFT JOIN dependencia ON contrato.idDependecia = dependencia.id
                LEFT JOIN objeto ON contrato.idObjeto = objeto.id
                LEFT JOIN obligacionesespecificas ON contrato.ObEspCon = obligacionesespecificas.id
                """

                # Ejecutar la consulta
                cursor.execute(consulta)
                registros = cursor.fetchall()

                # Crear un DataFrame de pandas con los resultados
                columnas = [
                    'ID Contrato', 'ID Versión', 'ID Cliente', 'Descripción Contrato', 'Vigencia', 'Terminación', 
                    'Autorización', 'Valor Contrato', 'Consecutivo', 'Nombre Cliente', 
                    'Tipo de Contrato', 'Cargo', 'Dependencia', 'Objeto', 'Obligación Específica'
                ]
                df = pd.DataFrame(registros, columns=columnas)

                # Asegurar que las fechas se formateen adecuadamente
                df['Vigencia'] = pd.to_datetime(df['Vigencia']).dt.strftime('%Y-%m-%d')
                df['Terminación'] = pd.to_datetime(df['Terminación']).dt.strftime('%Y-%m-%d')

                # Exportar el DataFrame a un archivo Excel en la ruta seleccionada
                df.to_excel(archivo_excel, index=False)

                # Ajustar el ancho de todas las columnas, incluida "Obligación Específica"
                ajustar_ancho_columnas(archivo_excel)

                print(f"Registros exportados a {archivo_excel} exitosamente.")
                messagebox.showinfo("Exportación exitosa", f"Registros exportados a {archivo_excel}")

        except Exception as e:
            print(f"Error al exportar a Excel: {e}")
            messagebox.showerror("Error", f"Error al exportar a Excel: {e}")

        finally:
            conexion.close()


    def ajustar_ancho_columnas(archivo_excel):
        # Cargar el archivo Excel recién creado
        libro = load_workbook(archivo_excel)
        hoja = libro.active

        # Iterar sobre las columnas y ajustar el ancho basado en el contenido
        for col in hoja.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)  # Obtener la letra de la columna

            for cell in col:
                try:
                    # Obtener la longitud máxima de los valores en la columna
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass

            # Ajustar el ancho de la columna (agregar un pequeño margen)
            hoja.column_dimensions[col_letter].width = max_length + 2

        # Guardar el archivo con los cambios
        libro.save(archivo_excel)

    def importar_desde_excel():
        try:
            # Abrir un diálogo para que el usuario seleccione el archivo Excel
            archivo_excel = filedialog.askopenfilename(
                filetypes=[("Archivos de Excel", "*.xlsx")],
                title="Seleccionar archivo de Excel"
            )
            
            if archivo_excel:  # Si el usuario selecciona un archivo
                # Leer el archivo Excel
                df = pd.read_excel(archivo_excel)

                # Conectar a la base de datos
                conexion = pymysql.connect(host='localhost', user='root', password='', db='contrataciones')
                cursor = conexion.cursor()

                # Recorrer el DataFrame y preparar la inserción de datos
                for index, row in df.iterrows():
                    # Obtener el ID del tipo de contrato a partir del nombre
                    cursor.execute("SELECT id FROM tipodecontrato WHERE nombreTipoContrato = %s", (row['Tipo de Contrato'],))
                    id_tipo_contrato = cursor.fetchone()[0]

                    # Obtener el ID del cargo a partir del nombre
                    cursor.execute("SELECT id FROM cargo WHERE nombreCargo = %s", (row['Cargo'],))
                    id_cargo = cursor.fetchone()[0]

                    # Obtener el ID de la dependencia a partir del nombre
                    cursor.execute("SELECT id FROM dependencia WHERE nombreDependencia = %s", (row['Dependencia'],))
                    id_dependencia = cursor.fetchone()[0]

                    # Obtener el ID del objeto a partir del nombre
                    cursor.execute("SELECT id FROM objeto WHERE nombreObjeto = %s", (row['Objeto'],))
                    id_objeto = cursor.fetchone()[0]

                    # Obtener el ID de la obligación específica a partir del nombre
                    cursor.execute("SELECT id FROM obligacionesespecificas WHERE nombre = %s", (row['Obligación Específica'],))
                    id_obligacion = cursor.fetchone()[0]

                    # Preparar la inserción de datos, excluyendo "Nombre Cliente"
                    insertar_contrato = """
                    INSERT INTO contrato (idVersion, idClientes, descripcionContrato, Vigencia, terminacion, autorizacionContratos, valorContrato, consecutivo, 
                                        idTipoContrato, idCargo, idDependecia, idObjeto, ObEspCon)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """
                    datos = (
                        row['ID Versión'],
                        row['ID Cliente'],  # Este se mantiene
                        row['Descripción Contrato'],
                        row['Vigencia'],
                        row['Terminación'],
                        row['Autorización'],
                        row['Valor Contrato'],
                        row['Consecutivo'],
                        id_tipo_contrato,  # Guardamos el ID del tipo de contrato
                        id_cargo,  # Guardamos el ID del cargo
                        id_dependencia,  # Guardamos el ID de la dependencia
                        id_objeto,  # Guardamos el ID del objeto
                        id_obligacion  # Guardamos el ID de la obligación específica
                    )

                    # Ejecutar la inserción en la tabla de contratos
                    cursor.execute(insertar_contrato, datos)

                # Confirmar los cambios
                conexion.commit()

                print("Datos importados exitosamente.")
                messagebox.showinfo("Importación exitosa", "Datos importados exitosamente desde el archivo Excel.")

        except Exception as e:
            print(f"Error al importar desde Excel: {e}")
            messagebox.showerror("Error", f"Error al importar desde Excel: {e}")

        finally:
            conexion.close()
    # Colocando cada botón en una columna diferente de la misma fila
    btnExportarWord = tk.Button(marco_botones, text="Exportar a Word", command=exportar_a_word)
    btnExportarWord.grid(row=0, column=0, padx=5, pady=5, sticky="ew")  # Primera columna

    btnExportarPDF = tk.Button(marco_botones, text="Exportar a PDF", command=exportar_a_pdf)
    btnExportarPDF.grid(row=0, column=1, padx=5, pady=5, sticky="ew")  # Segunda columna

    btnNuevaVersion = tk.Button(marco_botones, text="Nueva Versión", command=nueva_version)
    btnNuevaVersion.grid(row=0, column=2, padx=5, pady=5, sticky="ew")  # Tercera columna

    btnExportarExcel = tk.Button(marco_botones, text="Exportar a Excel", command=exportar_a_excel)
    btnExportarExcel.grid(row=0, column=3, padx=5, pady=5, sticky="ew")  # Cuarta columna

    btnImportarExcel = tk.Button(marco_botones, text="Importar desde Excel", command=importar_desde_excel)
    btnImportarExcel.grid(row=0, column=4, padx=5, pady=5, sticky="ew")  # Quinta columna

    # Expande cada botón para ocupar el ancho de su columna
    for i in range(5):
        marco_botones.grid_columnconfigure(i, weight=1)
    # Cargar datos en la tabla
    llenar_tabla()
    cargar_combo_data()
    actualizar_botones_paginacion()


############################################################### Pantalla contratistas ############################################################################
def mostrar_Contratista(frame):
    # Configuración de la conexión a la base de datos
    db = mysql.connector.connect(host="localhost", user="root", password="", database="contrataciones")
    cursor = db.cursor()

    # Crear un nuevo marco dentro del Toplevel para usar grid
    frame_principal = tk.Frame(frame)
    frame_principal.pack(fill='both', expand=True)

    # Configuración del marco principal
    marco = tk.LabelFrame(frame_principal, text="Formulario Contratistas")
    marco.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

    # Hacer que el marco sea redimensionable
    frame_principal.grid_rowconfigure(0, weight=1)
    frame_principal.grid_columnconfigure(0, weight=1)
    marco.grid_rowconfigure(11, weight=1)  # Permitir que la fila de la tabla se expanda
    marco.grid_columnconfigure(5, weight=1)  # Permitir que las columnas de widgets se expandan

    # Variables de paginación
    registros_por_pagina = 10
    pagina_actual = 0
    total_paginas = 0

    # Función para obtener opciones para los comboboxes
    def obtener_opciones(query):
        cursor.execute(query)
        return [(fila[0], fila[1]) for fila in cursor.fetchall()]

    # Consultas para llenar los comboboxes
    opciones_tipo_documento = obtener_opciones("SELECT id, numeroCedula FROM tipodedocumento")
    opciones_departamento = obtener_opciones("SELECT id, nombreDepartamento FROM departamentos")
    opciones_ciudad = obtener_opciones("SELECT id, nombreCiudad FROM ciudad")
    opciones_genero = obtener_opciones("SELECT id, nombreGenero FROM genero")
    opciones_rh = obtener_opciones("SELECT id, tipoRh FROM rh")
    opciones_arl = obtener_opciones("SELECT id, nombreARL FROM arl")
    opciones_eps = obtener_opciones("SELECT id, nombreEPS FROM eps")
    opciones_banco = obtener_opciones("SELECT id, nombreBanco FROM banco")
    opciones_estado = obtener_opciones("SELECT id, tipoEstado FROM estado")
    opciones_usuario_creador = obtener_opciones("SELECT id, numeroCedulaCliente FROM usuarios")

    # Variables de los campos de la tabla clientes
    id_cliente = tk.StringVar()
    numero_documento = tk.StringVar()
    primerNombre = tk.StringVar()
    segundoNombre = tk.StringVar()
    primerApellido = tk.StringVar()
    segundoApellido = tk.StringVar()
    tipoDocumento = tk.StringVar()
    departamentoExpedicion = tk.StringVar()
    ciudadExpedicion = tk.StringVar()
    fechaExpedicion = tk.StringVar()
    genero = tk.StringVar()
    fechaNacimento = tk.StringVar()
    Rh = tk.StringVar()
    dirrecion = tk.StringVar()
    correo = tk.StringVar()
    correoAdicional = tk.StringVar()
    celular = tk.StringVar()
    telefono = tk.StringVar()
    ARL = tk.StringVar()
    EPS = tk.StringVar()
    PAA = tk.StringVar()
    idBanco = tk.StringVar()
    numeroCuenta = tk.StringVar()
    CDP = tk.StringVar()
    fechaRegistro = tk.StringVar()
    ultimoAcceso = tk.StringVar()
    Estado = tk.StringVar()
    usuarioCreador = tk.StringVar()

    def validar_correo(correo):
        patron_correo = r"^[\w\.-]+@[\w\.-]+\.\w+$"
        return re.match(patron_correo, correo) is not None

    # Función para validar que todos los campos requeridos estén llenos
    def validar_campos():
        campos_faltantes = []
        if not numero_documento.get(): campos_faltantes.append("Número Documento")
        if not primerNombre.get(): campos_faltantes.append("Primer Nombre")
        if not segundoNombre.get(): campos_faltantes.append("Segundo Nombre")
        if not primerApellido.get(): campos_faltantes.append("Primer Apellido")
        if not segundoApellido.get(): campos_faltantes.append("Segundo Apellido")
        if not tipoDocumento.get(): campos_faltantes.append("Tipo Documento")
        if not departamentoExpedicion.get(): campos_faltantes.append("Departamento Expedición")
        if not ciudadExpedicion.get(): campos_faltantes.append("Ciudad Expedición")
        if not fechaExpedicion.get(): campos_faltantes.append("Fecha Expedición")
        if not genero.get(): campos_faltantes.append("Género")
        if not fechaNacimento.get(): campos_faltantes.append("Fecha Nacimiento")
        if not Rh.get(): campos_faltantes.append("RH")
        if not dirrecion.get(): campos_faltantes.append("Dirección")
        if not correo.get(): campos_faltantes.append("Correo")
        if not correoAdicional.get(): campos_faltantes.append("Correo Adicional")
        if not celular.get(): campos_faltantes.append("Celular")
        if not telefono.get(): campos_faltantes.append("Teléfono")
        if not ARL.get(): campos_faltantes.append("ARL")
        if not EPS.get(): campos_faltantes.append("EPS")
        if not PAA.get(): campos_faltantes.append("PAA")
        if not idBanco.get(): campos_faltantes.append("Banco")
        if not numeroCuenta.get(): campos_faltantes.append("Número Cuenta")
        if not CDP.get(): campos_faltantes.append("CDP")
        if not fechaRegistro.get(): campos_faltantes.append("Fecha Registro")
        if not ultimoAcceso.get(): campos_faltantes.append("Último Acceso")
        if not Estado.get(): campos_faltantes.append("Estado")
        if not usuarioCreador.get(): campos_faltantes.append("Usuario Creador")
        if not validar_correo(correo.get()): campos_faltantes.append("Correo válido")

        if campos_faltantes:
            lblMensaje.config(text=f"Faltan campos: {', '.join(campos_faltantes)}", fg="red", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            ajustar_texto_mensaje()
            return False
        return True

    def seleccionar(event):
        seleccion = tvClientes.selection()
        if seleccion:
            id = seleccion[0]
            valores = tvClientes.item(id, "values")

            # Asegurarnos de que el número de columnas coincida
            if len(valores) == len(tvClientes["columns"]):
                id_cliente.set(valores[0])
                numero_documento.set(valores[0])
                primerNombre.set(valores[1])
                segundoNombre.set(valores[2])
                primerApellido.set(valores[3])
                segundoApellido.set(valores[4])
                tipoDocumento.set(valores[5])
                departamentoExpedicion.set(valores[6])
                ciudadExpedicion.set(valores[7])
                fechaExpedicion.set(valores[8])
                genero.set(valores[9])
                fechaNacimento.set(valores[10])
                Rh.set(valores[11])
                dirrecion.set(valores[12])
                correo.set(valores[13])
                correoAdicional.set(valores[14])
                celular.set(valores[15])
                telefono.set(valores[16])
                ARL.set(valores[17])
                EPS.set(valores[18])
                PAA.set(valores[19])
                idBanco.set(valores[20])
                numeroCuenta.set(valores[21])
                CDP.set(valores[22])
                fechaRegistro.set(valores[23])
                ultimoAcceso.set(valores[24])
                Estado.set(valores[25])
                usuarioCreador.set(valores[26])

                # Bloquear el campo de Número de Documento al editar
                numero_documento_entry.config(state="readonly")
            else:
                lblMensaje.config(text="Error: Los datos seleccionados no tienen suficientes columnas.", fg="red", font=("Arial", 12, "bold"))
                ajustar_texto_mensaje()

    def limpiar_campos():
        id_cliente.set("")
        numero_documento.set("")
        primerNombre.set("")
        segundoNombre.set("")
        primerApellido.set("")
        segundoApellido.set("")
        tipoDocumento.set("")
        departamentoExpedicion.set("")
        ciudadExpedicion.set("")
        fechaExpedicion.set("")
        genero.set("")
        fechaNacimento.set("")
        Rh.set("")
        dirrecion.set("")
        correo.set("")
        correoAdicional.set("")
        celular.set("")
        telefono.set("")
        ARL.set("")
        EPS.set("")
        PAA.set("")
        idBanco.set("")
        numeroCuenta.set("")
        CDP.set("")
        fechaRegistro.set("")
        ultimoAcceso.set("")
        Estado.set("")
        usuarioCreador.set("")

        # Desbloquear el campo de Número de Documento
        numero_documento_entry.config(state="normal")

    # Función para filtrar las opciones en los comboboxes
    def filtrar_opciones(opciones, texto):
        texto_filtrado = texto.lower()
        return [txt for id, txt in opciones if texto_filtrado in txt.lower()]

    def actualizar_combobox(event, opciones, combobox):
        texto_actual = combobox.get()
        opciones_filtradas = filtrar_opciones(opciones, texto_actual)
        combobox['values'] = opciones_filtradas
        combobox.icursor(len(texto_actual))  # Mover el cursor al final del texto
        combobox.event_generate('<Down>')  # Mostrar lista desplegable

    # Configuración del marco de formulario
    numero_documento_entry = tk.Entry(marco, textvariable=numero_documento, width=20, validate="key", validatecommand=(frame.register(lambda s: s.isdigit()), "%P"))
    
    tipoDocumento_combobox = ttk.Combobox(marco, textvariable=tipoDocumento, values=[texto for id, texto in opciones_tipo_documento])
    tipoDocumento_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_tipo_documento, tipoDocumento_combobox))

    departamentoExpedicion_combobox = ttk.Combobox(marco, textvariable=departamentoExpedicion, values=[texto for id, texto in opciones_departamento])
    departamentoExpedicion_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_departamento, departamentoExpedicion_combobox))

    ciudadExpedicion_combobox = ttk.Combobox(marco, textvariable=ciudadExpedicion, values=[texto for id, texto in opciones_ciudad])
    ciudadExpedicion_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_ciudad, ciudadExpedicion_combobox))

    genero_combobox = ttk.Combobox(marco, textvariable=genero, values=[texto for id, texto in opciones_genero])
    genero_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_genero, genero_combobox))

    Rh_combobox = ttk.Combobox(marco, textvariable=Rh, values=[texto for id, texto in opciones_rh])
    Rh_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_rh, Rh_combobox))

    ARL_combobox = ttk.Combobox(marco, textvariable=ARL, values=[texto for id, texto in opciones_arl])
    ARL_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_arl, ARL_combobox))

    EPS_combobox = ttk.Combobox(marco, textvariable=EPS, values=[texto for id, texto in opciones_eps])
    EPS_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_eps, EPS_combobox))

    idBanco_combobox = ttk.Combobox(marco, textvariable=idBanco, values=[texto for id, texto in opciones_banco])
    idBanco_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_banco, idBanco_combobox))

    Estado_combobox = ttk.Combobox(marco, textvariable=Estado, values=[texto for id, texto in opciones_estado])
    Estado_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_estado, Estado_combobox))

    usuarioCreador_combobox = ttk.Combobox(marco, textvariable=usuarioCreador, values=[texto for id, texto in opciones_usuario_creador])
    usuarioCreador_combobox.bind('<KeyRelease>', lambda event: actualizar_combobox(event, opciones_usuario_creador, usuarioCreador_combobox))

    filas = [
        ("Número Documento", numero_documento, numero_documento_entry),
        ("Primer Nombre", primerNombre, tk.Entry(marco, textvariable=primerNombre, width=20)),
        ("Segundo Nombre", segundoNombre, tk.Entry(marco, textvariable=segundoNombre, width=20)),
        ("Primer Apellido", primerApellido, tk.Entry(marco, textvariable=primerApellido, width=20)),
        ("Segundo Apellido", segundoApellido, tk.Entry(marco, textvariable=segundoApellido, width=20)),
        ("Tipo Documento", tipoDocumento, tipoDocumento_combobox),
        ("Departamento Expedición", departamentoExpedicion, departamentoExpedicion_combobox),
        ("Ciudad Expedición", ciudadExpedicion, ciudadExpedicion_combobox),
        ("Fecha Expedición", fechaExpedicion, DateEntry(marco, textvariable=fechaExpedicion, width=20, background='darkblue', foreground='white', borderwidth=2, date_pattern='y-mm-dd')),
        ("Género", genero, genero_combobox),
        ("Fecha Nacimiento", fechaNacimento, DateEntry(marco, textvariable=fechaNacimento, width=20, background='darkblue', foreground='white', borderwidth=2, date_pattern='y-mm-dd')),
        ("RH", Rh, Rh_combobox),
        ("Dirección", dirrecion, tk.Entry(marco, textvariable=dirrecion, width=30)),
        ("Correo", correo, tk.Entry(marco, textvariable=correo, width=30)),
        ("Correo Adicional", correoAdicional, tk.Entry(marco, textvariable=correoAdicional, width=30)),
        ("Celular", celular, tk.Entry(marco, textvariable=celular, width=20)),
        ("Teléfono", telefono, tk.Entry(marco, textvariable=telefono, width=20)),
        ("ARL", ARL, ARL_combobox),
        ("EPS", EPS, EPS_combobox),
        ("PAA", PAA, tk.Entry(marco, textvariable=PAA, width=30)),
        ("Banco", idBanco, idBanco_combobox),
        ("Número Cuenta", numeroCuenta, tk.Entry(marco, textvariable=numeroCuenta, width=20)),
        ("CDP", CDP, tk.Entry(marco, textvariable=CDP, width=20)),
        ("Fecha Registro", fechaRegistro, DateEntry(marco, textvariable=fechaRegistro, width=20, background='darkblue', foreground='white', borderwidth=2, date_pattern='y-mm-dd')),
        ("Último Acceso", ultimoAcceso, DateEntry(marco, textvariable=ultimoAcceso, width=20, background='darkblue', foreground='white', borderwidth=2, date_pattern='y-mm-dd')),
        ("Estado", Estado, Estado_combobox),
        ("Usuario Creador", usuarioCreador, usuarioCreador_combobox)
    ]


    for i, (texto, var, widget) in enumerate(filas):
        columna = i % 3
        fila = i // 3
        tk.Label(marco, text=texto).grid(column=columna*2, row=fila, padx=2, pady=2, sticky="e")
        widget.grid(column=columna*2+1, row=fila, padx=2, pady=2, sticky="ew")

    lblMensaje = tk.Label(marco, text="Aquí van los mensajes", fg="green", font=("Arial", 12, "bold"))
    lblMensaje.grid(column=0, row=10, columnspan=6, padx=2, pady=2, sticky="ew")

    # Ajustar el texto del mensaje para adaptarse al tamaño de la ventana
    def ajustar_texto_mensaje():
        lblMensaje.config(wraplength=marco.winfo_width())

    marco.bind("<Configure>", lambda event: ajustar_texto_mensaje())

    # Scrollbars
    scrollbar_y = ttk.Scrollbar(marco, orient="vertical")
    scrollbar_y.grid(row=12, column=10, sticky="ns")
    scrollbar_x = ttk.Scrollbar(marco, orient="horizontal")
    scrollbar_x.grid(row=13, column=0, columnspan=9, sticky="ew")

    # Treeview con scrollbars
    tvClientes = ttk.Treeview(marco, yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set, selectmode='browse')
    tvClientes["columns"] = (
        "Primer Nombre", "Segundo Nombre", "Primer Apellido", "Segundo Apellido", "Tipo de Documento",
        "Departamento Expedicion", "Ciudad Expedición", "Fecha Expedición", "Género", "Fecha de Nacimento", "RH", "Dirreción", "Correo", "Correo Adicional",
        "Celular", "Teléfono", "ARL", "EPS", "PAA", "Banco", "Numero de Cuenta", "CDP", "fechaRegistro", "ultimoAcceso", "Estado", "usuarioCreador"
    )
    tvClientes.column("#0", width=0, stretch='no')
    tvClientes.column("primerNombre", width=100, anchor='center')

    for col in tvClientes["columns"]:
        tvClientes.column(col, width=125, anchor='center', stretch='yes')
        tvClientes.heading(col, text=col, anchor='center')

    tvClientes.grid(column=0, row=11, columnspan=10, padx=2, pady=2, sticky="nsew")
    scrollbar_y.config(command=tvClientes.yview)
    scrollbar_x.config(command=tvClientes.xview)
    marco.grid_columnconfigure(0, weight=1)
    
    marco.grid_rowconfigure(11, weight=1)
    marco.grid_rowconfigure(12, weight=0)

    query = ""

    def cargar_datos():
        nonlocal total_paginas, query
        # Consulta SQL con JOIN para obtener los textos correspondientes a las llaves foráneas
        query = """
        SELECT c.primerNombre, c.segundoNombre, c.primerApellido, c.segundoApellido,
            td.numeroCedula AS tipoDocumento, d.nombreDepartamento AS departamentoExpedicion,
            ci.nombreCiudad AS ciudadExpedicion, c.fechaExpedicion, g.nombreGenero AS genero,
            c.fechaNacimento, r.tipoRh AS Rh, c.dirrecion, c.correo, c.correoAdicional, c.celular, c.telefono,
            a.nombreARL AS ARL, e.nombreEPS AS EPS, c.PAA, b.nombreBanco AS idBanco,
            c.numeroCuenta, c.CDP, c.fechaRegistro, c.ultimoAcceso, es.tipoEstado AS Estado, u.numeroCedulaCliente AS usuarioCreador
        FROM clientes c
        JOIN tipodedocumento td ON c.tipoDocumento = td.id
        JOIN departamentos d ON c.departamentoExpedicion = d.id
        JOIN ciudad ci ON c.ciudadExpedicion = ci.id
        JOIN genero g ON c.genero = g.id
        JOIN rh r ON c.Rh = r.id
        JOIN arl a ON c.ARL = a.id
        JOIN eps e ON c.EPS = e.id
        JOIN banco b ON c.idBanco = b.id
        JOIN estado es ON c.Estado = es.id
        JOIN usuarios u ON c.usuarioCreador = u.id
        """
        cursor.execute(query)
        registros = cursor.fetchall()
        total_registros = len(registros)
        total_paginas = (total_registros + registros_por_pagina - 1) // registros_por_pagina
        mostrar_pagina(pagina_actual)

    def mostrar_pagina(pagina):
        tvClientes.delete(*tvClientes.get_children())
        offset = pagina * registros_por_pagina
        paginated_query = query + f" LIMIT {registros_por_pagina} OFFSET {offset}"
        cursor.execute(paginated_query)
        registros = cursor.fetchall()

        for registro in registros:
            tvClientes.insert("", "end", values=registro)
    def avanzar_pagina():
        nonlocal pagina_actual
        if pagina_actual < total_paginas - 1:
            pagina_actual += 1
            mostrar_pagina(pagina_actual)

    def retroceder_pagina():
        nonlocal pagina_actual
        if pagina_actual > 0:
            pagina_actual -= 1
            mostrar_pagina(pagina_actual)

    def obtener_id_por_texto(opciones, texto):
        for id, txt in opciones:
            if str(txt) == str(texto):  # Asegurarse de comparar los valores como cadenas
                return id
        return None

    def exportar_excel():
        archivo = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", ".xlsx"), ("All files", ".*")])
        if archivo:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Clientes"

            # Definir el estilo de borde
            borde = Border(
                left=Side(border_style="thin"),
                right=Side(border_style="thin"),
                top=Side(border_style="thin"),
                bottom=Side(border_style="thin")
            )

            # Escribir los encabezados de las columnas con borde
            for i, col in enumerate(tvClientes["columns"], 1):
                celda = ws.cell(row=1, column=i, value=col)
                celda.border = borde  # Aplicar borde a los encabezados

            # Escribir los valores del Treeview en el archivo Excel
            for i, item in enumerate(tvClientes.get_children(), 2):
                for j, valor in enumerate(tvClientes.item(item)['values'], 1):
                    ws.cell(row=i, column=j, value=valor)

            # Ajustar el ancho de las columnas basado en el contenido
            ajustar_ancho_columnas_excel(ws)

            # Guardar el archivo
            wb.save(archivo)
            lblMensaje.config(text=f"Datos exportados a {archivo}", fg="green", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")


    def ajustar_ancho_columnas_excel(hoja):
        # Iterar sobre las columnas y ajustar el ancho basado en el contenido
        for col in hoja.columns:
            max_length = 0
            col_letter = get_column_letter(col[0].column)  # Obtener la letra de la columna

            for cell in col:
                try:
                    # Obtener la longitud máxima de los valores en la columna
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                except:
                    pass

            # Ajustar el ancho de la columna (agregar un pequeño margen)
            hoja.column_dimensions[col_letter].width = max_length + 2

    # Función para importar datos desde un archivo Excel
    def importar_excel():
        archivo = filedialog.askopenfilename(defaultextension=".xlsx", filetypes=[("Excel files", ".xlsx"), ("All files", ".*")])
        if archivo:
            wb = openpyxl.load_workbook(archivo)
            ws = wb.active
            # Iterar sobre las filas y añadirlas a la base de datos
            for fila in ws.iter_rows(min_row=2, values_only=True):  # Saltar la fila de encabezado
                # Asumimos que el Excel tiene las columnas en el mismo orden que tu base de datos
                (
                    numero_documento, primerNombre, segundoNombre, primerApellido, segundoApellido, tipoDocumentoTexto,
                    departamentoExpedicionTexto, ciudadExpedicionTexto, fechaExpedicion, generoTexto, fechaNacimento,
                    RhTexto, dirrecion, correo, correoAdicional, celular, telefono, ARLTexto, EPSTexto, PAA,
                    bancoTexto, numeroCuenta, CDP, fechaRegistro, ultimoAcceso, estadoTexto, usuarioCreadorTexto
                ) = fila

                # Función para obtener el ID según el texto
                def obtener_id_por_texto(opciones, texto):
                    for id, txt in opciones:
                        if str(txt) == str(texto):
                            return id
                    return None

                # Obtener los IDs correspondientes a los textos de las llaves foráneas
                id_tipo_documento = obtener_id_por_texto(opciones_tipo_documento, tipoDocumentoTexto)
                id_departamento = obtener_id_por_texto(opciones_departamento, departamentoExpedicionTexto)
                id_ciudad = obtener_id_por_texto(opciones_ciudad, ciudadExpedicionTexto)
                id_genero = obtener_id_por_texto(opciones_genero, generoTexto)
                id_rh = obtener_id_por_texto(opciones_rh, RhTexto)
                id_arl = obtener_id_por_texto(opciones_arl, ARLTexto)
                id_eps = obtener_id_por_texto(opciones_eps, EPSTexto)
                id_banco = obtener_id_por_texto(opciones_banco, bancoTexto)
                id_estado = obtener_id_por_texto(opciones_estado, estadoTexto)
                id_usuario_creador = obtener_id_por_texto(opciones_usuario_creador, usuarioCreadorTexto)

                # Confirmar que todos los IDs se obtuvieron correctamente
                if None in [id_tipo_documento, id_departamento, id_ciudad, id_genero, id_rh, id_arl, id_eps, id_banco, id_estado, id_usuario_creador]:
                    lblMensaje.config(text=f"Error: No se pudieron obtener todos los IDs para la fila {fila}", fg="red", font=("Arial", 12, "bold"))
                    lblMensaje.grid(columnspan=6, sticky="ew")
                    continue  # Saltar esta fila si no se obtuvieron todos los IDs

                # Crear la tupla con los valores convertidos a IDs
                valores = (
                    numero_documento, primerNombre, segundoNombre, primerApellido, segundoApellido, id_tipo_documento,
                    id_departamento, id_ciudad, fechaExpedicion, id_genero, fechaNacimento, id_rh, dirrecion,
                    correo, correoAdicional, celular, telefono, id_arl, id_eps, PAA, id_banco, numeroCuenta, CDP,
                    fechaRegistro, ultimoAcceso, id_estado, id_usuario_creador
                )

                # Consulta SQL
                query = """
                INSERT INTO clientes (id, primerNombre, segundoNombre, primerApellido, segundoApellido, tipoDocumento,
                                    departamentoExpedicion, ciudadExpedicion, fechaExpedicion, genero, fechaNacimento, Rh, dirrecion,
                                    correo, correoAdicional, celular, telefono, ARL, EPS, PAA, idBanco, numeroCuenta, CDP, fechaRegistro, 
                                    ultimoAcceso, Estado, usuarioCreador)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                """
                try:
                    cursor.execute(query, valores)
                    db.commit()
                except mysql.connector.Error as err:
                    print(f"Error al insertar fila: {fila}")
                    print(err)
            lblMensaje.config(text=f"Datos importados desde {archivo}", fg="green", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            cargar_datos()

    def exportar_pdf():
        archivo = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", ".pdf"), ("All files", ".*")])
        if archivo:
            c = canvas.Canvas(archivo, pagesize=letter)
            width, height = letter

            y = height - 40
            for i, col in enumerate(tvClientes["columns"], 1):
                c.drawString(10 * i, y, col)

            y -= 20
            for item in tvClientes.get_children():
                for j, valor in enumerate(tvClientes.item(item)['values'], 1):
                    c.drawString(10 * j, y, str(valor))
                y -= 20
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            lblMensaje.config(text=f"Datos exportados a {archivo}", fg="green", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            
    def agregar_cliente():
        if not validar_campos():
            return

        # Obtener los IDs correspondientes a los textos seleccionados en los comboboxes
        id_tipo_documento = obtener_id_por_texto(opciones_tipo_documento, tipoDocumento.get())
        id_departamento = obtener_id_por_texto(opciones_departamento, departamentoExpedicion.get())
        id_ciudad = obtener_id_por_texto(opciones_ciudad, ciudadExpedicion.get())
        id_genero = obtener_id_por_texto(opciones_genero, genero.get())
        id_rh = obtener_id_por_texto(opciones_rh, Rh.get())
        id_arl = obtener_id_por_texto(opciones_arl, ARL.get())
        id_eps = obtener_id_por_texto(opciones_eps, EPS.get())
        id_banco = obtener_id_por_texto(opciones_banco, idBanco.get())
        id_estado = obtener_id_por_texto(opciones_estado, Estado.get())
        id_usuario_creador = obtener_id_por_texto(opciones_usuario_creador, usuarioCreador.get())

        if id_usuario_creador is None:
            lblMensaje.config(text="Error: No se ha seleccionado un usuario válido en 'Usuario Creador'.", fg="red", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            return

        # Generar la tupla de valores con exactamente 27 elementos
        valores = (
            numero_documento.get(), primerNombre.get(), segundoNombre.get(), primerApellido.get(), segundoApellido.get(),
            id_tipo_documento, id_departamento, id_ciudad, fechaExpedicion.get(), id_genero,
            fechaNacimento.get(), id_rh, dirrecion.get(), correo.get(), correoAdicional.get(), celular.get(), telefono.get(), id_arl, id_eps,
            PAA.get(), id_banco, numeroCuenta.get(), CDP.get(), fechaRegistro.get(), ultimoAcceso.get(), id_estado, id_usuario_creador,  # Aquí añadimos el valor extra
        )

        # Confirmar que la tupla tiene exactamente 27 valores
        if len(valores) != 27:
            print(f"Error: Número de valores no coincide con los placeholders. Valores: {len(valores)}")
            return

        # Consulta SQL
        query = """
        INSERT INTO clientes (id, primerNombre, segundoNombre, primerApellido, segundoApellido, tipoDocumento,
                            departamentoExpedicion, ciudadExpedicion, fechaExpedicion, genero, fechaNacimento, Rh, dirrecion,
                            correo, correoAdicional, celular, telefono, ARL, EPS, PAA, idBanco, numeroCuenta, CDP, fechaRegistro, ultimoAcceso, Estado, usuarioCreador)
        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
        """

        try:
            # Ejecutar la consulta
            cursor.execute(query, valores)
            db.commit()
            lblMensaje.config(text="Cliente agregado exitosamente", fg="green", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            cargar_datos()
        except mysql.connector.Error as err:
            # Imprimir el error en la consola
            print("Error durante la ejecución de la consulta SQL:", err)
            print("Consulta SQL:", query)
            print("Valores:", valores)

    def actualizar_cliente():
        if not validar_campos():
            return

        id_tipo_documento = obtener_id_por_texto(opciones_tipo_documento, tipoDocumento.get())
        id_departamento = obtener_id_por_texto(opciones_departamento, departamentoExpedicion.get())
        id_ciudad = obtener_id_por_texto(opciones_ciudad, ciudadExpedicion.get())
        id_genero = obtener_id_por_texto(opciones_genero, genero.get())
        id_rh = obtener_id_por_texto(opciones_rh, Rh.get())
        id_arl = obtener_id_por_texto(opciones_arl, ARL.get())
        id_eps = obtener_id_por_texto(opciones_eps, EPS.get())
        id_banco = obtener_id_por_texto(opciones_banco, idBanco.get())
        id_estado = obtener_id_por_texto(opciones_estado, Estado.get())
        id_usuario_creador = obtener_id_por_texto(opciones_usuario_creador, usuarioCreador.get())

        if id_usuario_creador is None:
            lblMensaje.config(text="Error: No se ha seleccionado un usuario válido en 'Usuario Creador'.", fg="red", font=("Arial", 12, "bold"))
            lblMensaje.grid(columnspan=6, sticky="ew")
            return

        query = """
        UPDATE clientes SET primerNombre=%s, segundoNombre=%s, primerApellido=%s, segundoApellido=%s, tipoDocumento=%s,
                            departamentoExpedicion=%s, ciudadExpedicion=%s, fechaExpedicion=%s, genero=%s, fechaNacimento=%s,
                            Rh=%s, dirrecion=%s, correo=%s, correoAdicional=%s, celular=%s, telefono=%s, ARL=%s, EPS=%s, PAA=%s, idBanco=%s,
                            numeroCuenta=%s, CDP=%s, fechaRegistro=%s, ultimoAcceso=%s, Estado=%s, usuarioCreador=%s
        WHERE id=%s
        """
        valores = (
            primerNombre.get(), segundoNombre.get(), primerApellido.get(), segundoApellido.get(), id_tipo_documento,
            id_departamento, id_ciudad, fechaExpedicion.get(), id_genero, fechaNacimento.get(),
            id_rh, dirrecion.get(), correo.get(), correoAdicional.get(), celular.get(), telefono.get(), id_arl, id_eps, PAA.get(),
            id_banco, numeroCuenta.get(), CDP.get(), fechaRegistro.get(), ultimoAcceso.get(), id_estado, id_usuario_creador,
            id_cliente.get()
        )
        cursor.execute(query, valores)
        db.commit()
        lblMensaje.config(text="Cliente actualizado exitosamente", fg="green", font=("Arial", 12, "bold"))
        lblMensaje.grid(columnspan=6, sticky="ew")
        cargar_datos()

    def eliminar_cliente():
        query = "DELETE FROM clientes WHERE id=%s"
        cursor.execute(query, (id_cliente.get(),))
        db.commit()
        lblMensaje.config(text="Cliente eliminado exitosamente", fg="green", font=("Arial", 12, "bold"))
        lblMensaje.grid(columnspan=6, sticky="ew")
        cargar_datos()

    # Botones de acción
    btnAgregar = ctk.CTkButton(marco, text="Agregar", fg_color=COLOR_BOTON, command=agregar_cliente)
    btnAgregar.grid(column=0, row=12, padx=5, pady=5)
    btnActualizar = ctk.CTkButton(marco, text="Actualizar", fg_color=COLOR_BOTON,command=actualizar_cliente)
    btnActualizar.grid(column=1, row=12, padx=5, pady=5)
    btnEliminar = ctk.CTkButton(marco, text="Eliminar", fg_color=COLOR_BOTON,command=eliminar_cliente)
    btnEliminar.grid(column=2, row=12, padx=5, pady=5)
    btnLimpiar = ctk.CTkButton(marco, text="Limpiar",fg_color=COLOR_BOTON, command=limpiar_campos)
    btnLimpiar.grid(column=3, row=12, padx=5, pady=5)
    btnExportarExcel = ctk.CTkButton(marco, text="Exportar Excel",fg_color=COLOR_BOTON, command=exportar_excel)
    btnExportarExcel.grid(column=4, row=12, padx=5, pady=5)
    btnImportarExcel = ctk.CTkButton(marco, text="Importar Excel",fg_color=COLOR_BOTON, command=importar_excel)
    btnImportarExcel.grid(column=5, row=12, padx=5, pady=5)
    btnExportarPDF = ctk.CTkButton(marco, text="Exportar PDF",fg_color=COLOR_BOTON, command=exportar_pdf)
    btnExportarPDF.grid(column=6, row=12, padx=5, pady=5)
    print("Botones creados correctamente")
    
    
    btnAgregar.lift()
    btnActualizar.lift()
    btnEliminar.lift()
    btnLimpiar.lift()
    btnExportarExcel.lift()
    btnImportarExcel.lift()
    btnExportarPDF.lift()

    cargar_datos()

############################################################## Pantalla usuarios ######################################################################
def mostrar_usuario(frame, db):
    # Variables
    numero_cedula_cliente = ctk.StringVar()
    correo_cli = ctk.StringVar()
    contrasena_cli = ctk.StringVar()
    accion_usuario = ctk.StringVar(value="Agregar")
    usuario_id = None
    
    def seleccionar_usuario(event):
        nonlocal usuario_id
        seleccion = tabla.selection()
        if seleccion:
            valores = tabla.item(seleccion[0], "values")
            if valores:
                usuario_id = valores[0]
                numero_cedula_cliente.set(valores[1])
                correo_cli.set(valores[2])
                contrasena_cli.set("******")  # No mostrar la contraseña real

    def llenar_tabla():
        tabla.delete(*tabla.get_children())
        for fila in db.obtener_datos("SELECT id, numeroCedulaCliente, correoCli FROM usuarios"):
            tabla.insert("", "end", values=fila)

    def ejecutar_accion():
        accion = accion_usuario.get().strip()
        if not numero_cedula_cliente.get() or not correo_cli.get() or not contrasena_cli.get():
            return

        if accion == "Agregar":
            hashed_password = bcrypt.hashpw(contrasena_cli.get().encode('utf-8'), bcrypt.gensalt())
            db.ejecutar_consulta("INSERT INTO usuarios (numeroCedulaCliente, idRol, correoCli, contraseñaCli) VALUES (%s, 1, %s, %s)",
                                 (numero_cedula_cliente.get(), correo_cli.get(), hashed_password))
        elif accion == "Modificar" and usuario_id:
            hashed_password = bcrypt.hashpw(contrasena_cli.get().encode('utf-8'), bcrypt.gensalt())
            db.ejecutar_consulta("UPDATE usuarios SET numeroCedulaCliente=%s, idRol=1, correoCli=%s, contraseñaCli=%s WHERE id=%s",
                                 (numero_cedula_cliente.get(), correo_cli.get(), hashed_password, usuario_id))
        elif accion == "Eliminar" and usuario_id:
            db.ejecutar_consulta("DELETE FROM usuarios WHERE id=%s", (usuario_id,))

        llenar_tabla()

    # Interfaz gráfica
    marco_usuario = ctk.CTkFrame(frame, fg_color="transparent")
    marco_usuario.pack(fill="both", expand=True, padx=20, pady=20)

    # Campos de entrada
    ctk.CTkLabel(marco_usuario, text="Cédula Usuario", text_color=COLOR_TEXTO).grid(row=0, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_usuario, textvariable=numero_cedula_cliente, fg_color="white", text_color=COLOR_TEXTO).grid(row=0, column=1, padx=5, pady=5)
    
    ctk.CTkLabel(marco_usuario, text="Correo", text_color=COLOR_TEXTO).grid(row=1, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_usuario, textvariable=correo_cli, fg_color="white", text_color=COLOR_TEXTO).grid(row=1, column=1, padx=5, pady=5)
    
    ctk.CTkLabel(marco_usuario, text="Contraseña", text_color=COLOR_TEXTO).grid(row=2, column=0, padx=5, pady=5)
    ctk.CTkEntry(marco_usuario, textvariable=contrasena_cli, show="*", fg_color="white", text_color=COLOR_TEXTO).grid(row=2, column=1, padx=5, pady=5)
    
    ctk.CTkLabel(marco_usuario, text="Acción", text_color=COLOR_TEXTO).grid(row=3, column=0, padx=5, pady=5)
    ctk.CTkComboBox(marco_usuario, variable=ejecutar_accion, fg_color="white", values=["Agregar", "Modificar", "Eliminar"]).grid(row=3, column=1, padx=5, pady=5)
    
    ctk.CTkButton(marco_usuario, text="Ejecutar", command=ejecutar_accion,fg_color=COLOR_BOTON).grid(row=4, column=0, columnspan=2, pady=10)
    
    # Tabla de usuarios
    frame_tabla = ctk.CTkFrame(marco_usuario)
    frame_tabla.grid(row=5, column=0, columnspan=2, sticky="nsew", padx=5, pady=5)
    
    columnas = ("ID", "Cédula", "Correo")
    tabla = ttk.Treeview(frame_tabla, columns=columnas, show="headings", height=10)
    tabla.heading("Cédula", text="Cédula Cliente")
    tabla.heading("Correo", text="Correo")
    tabla.column("ID", width=0, stretch=False)  # Ocultar ID
    tabla.column("Cédula", width=150)
    tabla.column("Correo", width=200)
    tabla.pack(side="left", fill="both", expand=True)
    tabla.bind("<<TreeviewSelect>>", seleccionar_usuario)

    scrollbar = ttk.Scrollbar(frame_tabla, orient="vertical", command=tabla.yview)
    tabla.configure(yscrollcommand=scrollbar.set)
    scrollbar.pack(side="right", fill="y")

    llenar_tabla()

############################################## Pantalla Bienvenida ###################################################
def pantalla_bienvenida():
    global usuario_iniciado
    for widget in pantalla.winfo_children():
        widget.destroy()
    
    if not usuario_iniciado:
        messagebox.showinfo(title="Acceso Denegado", message="Inicie sesión primero")
        return

    bienvenida = ctk.CTkToplevel()
    bienvenida.title("Vista de administrador ~ CENIGRAF")
    bienvenida.state("zoomed")  # Se ajusta a un tamaño manejable
    bienvenida.configure(fg_color="#AFFFAF")

    # Cargar fuente Zurich
    fuente_zurich = ctk.CTkFont(family="Zurich.ttf", weight="bold", size=14)

    # Barra superior
    top_frame = ctk.CTkFrame(bienvenida, fg_color="#4CD14C", height=50, corner_radius=0)
    top_frame.pack(side="top", fill="x")

    cerrar_sesion_button = ctk.CTkButton(
        top_frame, text="Cerrar Sesión", command=lambda: cerrar_sesion(bienvenida),
        font=fuente_zurich, fg_color="#FF0022", hover_color="#9E0015",
        corner_radius=10, border_width=2, border_color="#9E0015"
    )
    cerrar_sesion_button.pack(side="right", padx=10, pady=10)

    # Contenedor principal
    main_frame = ctk.CTkFrame(bienvenida, fg_color="transparent")
    main_frame.pack(fill="both", expand=True)

    # Menú lateral fijo
    menu_frame = ctk.CTkFrame(
        main_frame, fg_color="#5DFF5D", width=220, corner_radius=15,
        border_width=2, border_color="white"
    )
    menu_frame.pack(side="left", fill="y", padx=10, pady=10)

    # Contenedor de contenido
    contenido_frame = ctk.CTkFrame(main_frame, fg_color="#AFFFAF", corner_radius=15)
    contenido_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

    # Opciones del menú
    pantallas = [
        "Pantalla EPS y ARL", "Pantalla Bancos", "Pantalla Ciudad y Departamento", 
        "Pantalla Cargo","Pantalla Tipo de contrato", "Pantalla Áreas", "Pantalla Certificado",
        "Pantalla Contratistas", "Pantalla Usuarios"
    ]

    for texto in pantallas:
        boton = ctk.CTkButton(
            menu_frame, text=texto, command=lambda t=texto: abrir_pantalla(contenido_frame, t),
            font=fuente_zurich, fg_color="#058A15", hover_color="#035E0E",
            corner_radius=10, border_width=2, border_color="#058A15"
        )
        boton.pack(fill="x", padx=10, pady=15)

######################################################################### Pantalla de inicio ########################################################################################
def mostrar_pantalla_inicial():
    global pantalla, image
    if pantalla is None:
        pantalla = ctk.CTk()
        pantalla.geometry("900x600")
        pantalla.resizable(False,False)
        pantalla.title("Contratistas CENIGRAF")
        pantalla.configure(fg_color="white")

    # Limpiar la pantalla
    for widget in pantalla.winfo_children():
        widget.destroy()
    
    # Intentar cargar el icono de la aplicación
    try:
        pantalla.iconbitmap("logo.ico")
    except:
        print("Advertencia: No se pudo cargar el icono 'logo.ico'")

    # Cargar y ajustar la imagen de fondo
    try:
        image = Image.open("fondo.png")
        image = image.resize((900, 600))  # Ajustar tamaño sin deformar
        image = ctk.CTkImage(image, size=(900, 600))
        fondo_label = ctk.CTkLabel(pantalla, image=image, text="")  # Fondo sin texto
        fondo_label.place(relwidth=1, relheight=1)
    except:
        print("Advertencia: No se pudo cargar la imagen.")

    # Fuente personalizada
    fuente_titulo = ctk.CTkFont(family="Zurich", size=32, weight="bold")
    fuente_botones = ctk.CTkFont(family="Zurich", size=22)

    # Título de acceso
    acceso_label = ctk.CTkLabel(pantalla, text="Acceso al sistema", fg_color="white", text_color=COLOR_BOTON, font=fuente_titulo)
    acceso_label.place(relx=0.5, rely=0.25, anchor="center")

    # Contenedor de botones con padding y diseño limpio
    botones_frame = ctk.CTkFrame(pantalla, fg_color="white", corner_radius=0)
    botones_frame.place(relx=0.5, rely=0.8, anchor="center")

    # Botón "Iniciar Sesión"
    btn_inicio = ctk.CTkButton(
        botones_frame, text="Iniciar Sesión", command=inicio_sesion,
        width=200, height=50, font=fuente_botones,
        fg_color=COLOR_BOTON, hover_color=COLOR_BOTON_HOVER, corner_radius=10
    )
    btn_inicio.pack(pady=15, padx=20)

    # Botón "Restaurar Contraseña"
    btn_restaurar = ctk.CTkButton(
        botones_frame, text="Restaurar Contraseña", command=mostrar_pantalla_restaurar,
        width=200, height=50, font=fuente_botones,
        fg_color=COLOR_BOTON, hover_color=COLOR_BOTON_HOVER, corner_radius=10
    )
    btn_restaurar.pack(pady=15, padx=20)

    pantalla.mainloop()

############################################################################## Validar inicio de sesion ##############################################################################
def validacion_datos(pantalla_sesion):
    global usuario_iniciado
    usuario = usuario_entry.get().strip()
    contrasena = contra_entry.get().strip()

    # Conectar a la base de datos
    bd = pymysql.connect(
        host="localhost",
        user="root",
        password="",
        db="contrataciones"
    )
    fcursor = bd.cursor()

    if not usuario or not contrasena:
        messagebox.showinfo(title="Error", message="Por favor, ingrese ambos campos")
        return

    # Consultar la contraseña en la base de datos
    sql = "SELECT contraseñaCli, idRol FROM usuarios WHERE correoCli=%s"
    fcursor.execute(sql, (usuario,))
    resultado = fcursor.fetchone()

    if resultado:
        hashed_password, rol = resultado  # Extraemos la contraseña y el rol
        if bcrypt.checkpw(contrasena.encode('utf-8'), hashed_password.encode('utf-8')):
            # Dependiendo del rol, abre una ventana diferente
            usuario_iniciado=True
            pantalla_sesion.withdraw()
            if rol == 1:
                pantalla_bienvenida()
            elif rol == 2:
                pantalla_bienvenida()
                
            else:
                messagebox.showerror("Error", "Rol desconocido")
        else:
            messagebox.showerror("Error", "Contraseña incorrecta")
    else:
        messagebox.showerror("Error", "Usuario no encontrado")
    fcursor.close()
    bd.close()

###################################################################################### Inicio de Sesion ##############################################################################
def inicio_sesion(): 
    global usuario_entry, contra_entry, image
    for widget in pantalla.winfo_children():
        widget.destroy()  # Limpiar la pantalla antes de mostrar la nueva ventana

    # Cargar fondo de pantalla
    try:
        image = Image.open("fondoLogin.png")
        image = image.resize((900, 600))  # Ajustar tamaño sin deformar
        image = ctk.CTkImage(image, size=(900, 600))
        fondo_label = ctk.CTkLabel(pantalla, image=image, text="")
        fondo_label.place(relwidth=1, relheight=1)
    except:
        print("Advertencia: No se pudo cargar la imagen de fondo.")

    fuente_titulos = ctk.CTkFont(family="Zurich.ttf", weight="bold", size=32)
    fuente = ctk.CTkFont(family="Zurich", size=24)
    fuente_botones = ctk.CTkFont(family="Zurich", size=22)

    # Título de acceso
    acceso_label = ctk.CTkLabel(pantalla, text="Iniciar Sesión", fg_color="white", text_color=COLOR_BOTON, font=fuente_titulos)
    acceso_label.place(relx=0.5, rely=0.2, anchor="center")

    # Contenedor principal con bordes redondeados
    container = ctk.CTkFrame(pantalla, fg_color="white", corner_radius=20)
    container.place(relx=0.5, rely=0.55, anchor="center")

    # Campos de entrada
    ctk.CTkLabel(container, text="Correo electrónico", text_color="#446344", font=fuente).pack(anchor="w", pady=5, padx=20)
    usuario_entry = ctk.CTkEntry(container, width=300, font=fuente, fg_color="white", text_color="black", border_width=2)
    usuario_entry.pack(pady=10, padx=20)

    ctk.CTkLabel(container, text="Contraseña", text_color="#446344", font=fuente).pack(anchor="w", pady=5, padx=20)
    contra_entry = ctk.CTkEntry(container, show="*", width=300, font=fuente, fg_color="white", text_color="black", border_width=2)
    contra_entry.pack(pady=10, padx=20)

    # Contenedor de botones con bordes redondeados
    botones_frame = ctk.CTkFrame(container, fg_color="white", corner_radius=10)  
    botones_frame.pack(pady=25)

    ctk.CTkButton(botones_frame, text="Iniciar Sesión", command=lambda: validacion_datos(pantalla),
                  width=120, height=40, font=fuente_botones, fg_color=COLOR_BOTON, hover_color=COLOR_BOTON_HOVER).grid(row=0, column=0, padx=15)

    ctk.CTkButton(botones_frame, text="Regresar", command=mostrar_pantalla_inicial, 
                  width=120, height=40, font=fuente_botones, fg_color=COLOR_BOTON, hover_color=COLOR_BOTON_HOVER).grid(row=0, column=1, padx=15)
################################################################################ Restaurar contraseña ##############################################################################

def validar_datos():
    cedula = entry_cedula.get().strip()
    correo = entry_correo.get().strip()

    if not cedula or not correo:
        messagebox.showerror("Error", "Por favor, complete ambos campos.")
        return

    bd = pymysql.connect(
            host="localhost",
            user="root",
            password="",
            db="contrataciones"
        )
    fcursor = bd.cursor()
    sql="SELECT id FROM usuarios WHERE numeroCedulaCliente =%s AND correoCli =%s"
    fcursor.execute(sql,(cedula,correo))
    usuario = fcursor.fetchone()
    if usuario:
        mostrar_campo_contraseña(usuario[0])
    else:
        messagebox.showerror("Error", "No se encontró un usuario con los datos ingresados.")
    fcursor.close()

def mostrar_campo_contraseña(user_id):
    global entry_nueva_contraseña, boton_cambiar

    if "entry_nueva_contraseña" in globals():
        return  # Evitar duplicar el campo de contraseña
    
    fuente = ctk.CTkFont(family="Zurich", size=24)
    fuente_botones = ctk.CTkFont(family="Zurich", size=22)

    # Campo de nueva contraseña
    ctk.CTkLabel(container, text="Nueva Contraseña", text_color="#446344", font=fuente).grid(row=5, column=0, sticky="w", pady=5, padx=10)
    entry_nueva_contraseña = ctk.CTkEntry(container, width=300, font=fuente, show="*", fg_color="white", text_color="black")
    entry_nueva_contraseña.grid(row=6, column=0, pady=5, padx=10)

    # Botón para actualizar la contraseña
    boton_cambiar = ctk.CTkButton(container, text="Actualizar Contraseña", command=lambda: actualizar_contraseña(user_id),
                                  width=200, font=fuente_botones, fg_color="green")
    boton_cambiar.grid(row=7, column=0, pady=5)

def actualizar_contraseña(user_id):
    nueva_contraseña = entry_nueva_contraseña.get().strip()

    if not nueva_contraseña:
        messagebox.showerror("Error", "La nueva contraseña no puede estar vacía.")
        return

    # Encriptar la contraseña con bcrypt
    hashed_password = bcrypt.hashpw(nueva_contraseña.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    try:
        bd = pymysql.connect(host="localhost", user="root", password="", db="contrataciones")
        fcursor = bd.cursor()
        sql = "UPDATE usuarios SET contraseñaCli = %s WHERE id = %s"
        fcursor.execute(sql, (hashed_password, user_id))
        bd.commit()
        messagebox.showinfo("Éxito", "Contraseña actualizada correctamente.")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo actualizar la contraseña.\n{str(e)}")
    finally:
        fcursor.close()
        bd.close()

    inicio_sesion()

def mostrar_pantalla_restaurar():
    global pantalla, entry_cedula, entry_correo, container

    for widget in pantalla.winfo_children():
        widget.destroy()  # Limpiar la pantalla antes de cargar la nueva vista

    # Fondo de pantalla
    try:
        image = Image.open("fondoLogin.png")
        image = image.resize((900, 600))  # Ajustar tamaño sin deformar
        image = ctk.CTkImage(image, size=(900, 600))
        fondo_label = ctk.CTkLabel(pantalla, image=image, text="")
        fondo_label.place(relwidth=1, relheight=1)
    except:
        print("Advertencia: No se pudo cargar la imagen de fondo.")

    fuente_titulos = ctk.CTkFont(family="Zurich.ttf", weight="bold", size=32)
    fuente = ctk.CTkFont(family="Zurich", size=24)
    fuente_botones = ctk.CTkFont(family="Zurich", size=22)

    # Título
    acceso_label = ctk.CTkLabel(pantalla, text="Restaurar Contraseña", fg_color="white", text_color=COLOR_BOTON, font=fuente_titulos)
    acceso_label.place(relx=0.5, rely=0.2, anchor="center")

    # Contenedor para los campos
    container = ctk.CTkFrame(pantalla, fg_color="white", corner_radius=20)
    container.place(relx=0.5, rely=0.55, anchor="center")

    # Etiquetas y entradas
    ctk.CTkLabel(container, text="Número de Cédula", text_color="#446344", font=fuente).grid(row=0, column=0, sticky="w", pady=5, padx=10)
    entry_cedula = ctk.CTkEntry(container, width=300, font=fuente, fg_color="white", text_color="black")
    entry_cedula.grid(row=1, column=0, pady=5, padx=10)

    ctk.CTkLabel(container, text="Correo Electrónico", text_color="#446344", font=fuente).grid(row=2, column=0, sticky="w", pady=5, padx=10)
    entry_correo = ctk.CTkEntry(container, width=300, font=fuente, fg_color="white", text_color="black")
    entry_correo.grid(row=3, column=0, pady=5, padx=10)

    # Contenedor para botones
    botones_frame = ctk.CTkFrame(container, fg_color="white")
    botones_frame.grid(row=4, column=0, pady=15)

    ctk.CTkButton(botones_frame, text="Validar", width=120, height=40, font=fuente_botones, fg_color=COLOR_BOTON, command=validar_datos).grid(row=0, column=0, padx=10, pady=15)
    ctk.CTkButton(botones_frame, text="Regresar", command=mostrar_pantalla_inicial, width=120, height=40, font=fuente_botones, fg_color=COLOR_BOTON).grid(row=0, column=1, padx=10, pady=15)
# Función para cerrar sesión
def cerrar_sesion(ventana):
    global usuario_iniciado,pantalla
    usuario_iniciado = False
    ventana.destroy()
    pantalla=None
    inicio_sesion()

# Iniciar la aplicación
mostrar_pantalla_inicial()